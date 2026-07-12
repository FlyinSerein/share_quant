from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

from build_docx_research_report import (
    BASE,
    IMAGES,
    OUT,
    TABLES,
    add_bullets,
    add_cover,
    add_figure,
    add_heading,
    add_para,
    add_table,
    num,
    pct,
    setup_document,
)


def add_docx_report() -> Path:
    single_top = pd.read_csv(TABLES / "single_factor_top_metrics.csv")
    single_ls = pd.read_csv(TABLES / "single_factor_long_short_metrics.csv")
    layer_metrics = pd.read_csv(TABLES / "layer_metrics.csv")
    long_short = pd.read_csv(TABLES / "long_short_metrics.csv")
    coverage = pd.read_csv(TABLES / "composite_scores_coverage.csv")
    ic = pd.read_csv(TABLES / "composite_ic_summary.csv")
    layer_summary = pd.read_csv(TABLES / "layer_summary.csv")
    weights = pd.read_csv(TABLES / "composite_factor_weights.csv")

    doc = Document()
    setup_document(doc)
    add_cover(doc)

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "本报告在已有单因子构建、检测与回测结果基础上，进一步研究多因子合成方法的表现。研究对象包括两个组合因子：等权合成因子 Composite_Equal，以及基于过去 12 期 RankIC 动态加权的 Composite_RollingRankIC。",
    )
    add_para(
        doc,
        "回测显示，滚动 RankIC 加权多因子显著优于简单等权合成，但在当前样本期内仍低于 PE_TTM、Volatility、Dividend_Yield 等表现最强的单因子。该结果说明多因子合成的主要价值不是追求样本内最高收益，而是降低单一风格暴露和单因子失效风险。",
    )

    add_heading(doc, "1. 研究背景与目标", 1)
    add_para(
        doc,
        "单因子研究已经完成因子构造、Top20% 组合回测、IC 诊断、十分组收益和风险暴露分析。本阶段的目标是在该基础上检验多因子合成是否能改善排序能力与组合稳健性，并解释其与强单因子的相对表现。",
    )
    add_para(doc, "本报告重点回答以下问题：")
    add_bullets(
        doc,
        [
            "等权多因子是否能够直接提升组合收益和横截面排序能力。",
            "滚动 RankIC 加权是否能降低弱因子拖累并改善多因子表现。",
            "多因子未跑赢部分强单因子时，是否仍具有合理的金融意义和后续研究价值。",
        ],
    )

    add_heading(doc, "2. 回测口径", 1)
    add_para(doc, "本次多因子研究沿用单因子研究的数据层和交易假设，核心口径如下：")
    add_table(
        doc,
        ["项目", "设置"],
        [
            ["评价期", "2022-01-01 至 2026-07-07"],
            ["调仓频率", "月度调仓，月末生成信号，下一交易日成交"],
            ["股票池", "已上市、非停牌、非 ST；每期至少 6 个有效单因子得分"],
            ["信号处理", "方向统一、缩尾、z-score 标准化，并做行业和市值中性化"],
            ["分层方式", "按综合得分分为 D1 至 D10，D10 为最高分组"],
            ["交易成本", "单边 10bp，在组合换手时扣除"],
            ["基准", "中证全指 000985.CSI"],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "3. 多因子构造方法", 1)
    add_para(
        doc,
        "本次研究构造两个多因子版本。Composite_Equal 对每只股票当期可用的中性化单因子得分做等权平均。Composite_RollingRankIC 使用各单因子过去 12 期 RankIC 均值作为权重，负权重截断为 0；当历史期数不足或所有权重为 0 时退回等权。",
    )
    add_table(
        doc,
        ["组合因子", "构造方式", "研究目的"],
        [
            ["Composite_Equal", "可用中性化单因子得分等权平均", "检验简单平均是否有效"],
            ["Composite_RollingRankIC", "过去 12 期 RankIC 均值加权，负权重截断为 0", "根据历史有效性动态调整因子权重"],
        ],
        [1.75, 3.05, 1.7],
    )
    active_weights = (
        weights[(weights["composite_factor"].eq("Composite_RollingRankIC")) & (weights["weight_source"].eq("rolling_rankic"))]
        .groupby("source_factor")["weight"]
        .mean()
        .sort_values(ascending=False)
    )
    add_table(doc, ["源因子", "平均权重"], [[idx, pct(value)] for idx, value in active_weights.head(11).items()], [2.8, 1.2], numeric_cols={1})
    add_para(doc, "权重主要集中在 Volatility、Dividend_Yield 和 PE_TTM，说明滚动 RankIC 机制识别出的有效因子与单因子回测结果基本一致。")

    add_heading(doc, "4. 样本覆盖与 IC 诊断", 1)
    coverage_rows = [
        [
            row["factor"],
            f"{int(row['signal_count'])}",
            str(int(row["first_signal"])),
            str(int(row["last_signal"])),
            f"{row['average_valid_stocks']:.0f}",
            f"{row['average_available_factor_count']:.2f}",
        ]
        for _, row in coverage.iterrows()
    ]
    add_table(doc, ["组合因子", "调仓次数", "首个信号", "最后信号", "平均有效股票数", "平均可用因子数"], coverage_rows, [1.7, 0.75, 0.95, 0.95, 1.1, 1.05], numeric_cols={1, 2, 3, 4, 5})
    ic_rows = [
        [row["composite_factor"], pct(row["ic_mean"]), pct(row["ic_win_rate"]), pct(row["rank_ic_mean"]), pct(row["rank_ic_win_rate"]), num(row["rank_ic_annual_icir"])]
        for _, row in ic.sort_values("rank_ic_mean", ascending=False).iterrows()
    ]
    add_table(doc, ["组合因子", "IC均值", "IC胜率", "RankIC均值", "RankIC胜率", "年化RankIC IR"], ic_rows, [1.85, 0.8, 0.8, 0.9, 0.9, 1.0], numeric_cols={1, 2, 3, 4, 5})
    add_para(doc, "Composite_RollingRankIC 的 RankIC 均值为 5.69%，RankIC 胜率为 68.52%，明显高于等权合成，说明动态加权改善了横截面排序能力。")

    add_heading(doc, "5. 多因子分层回测结果", 1)
    add_heading(doc, "5.1 Top 层组合表现", 2)
    top = layer_metrics[layer_metrics["bucket"].eq(10)].sort_values("annual_return", ascending=False)
    top_rows = [
        [row["composite_factor"], pct(row["annual_return"]), pct(row["annual_volatility"]), pct(row["max_drawdown"]), num(row["sharpe"]), pct(row["excess_annual_return"]), pct(row["average_monthly_turnover"])]
        for _, row in top.iterrows()
    ]
    add_table(doc, ["组合因子", "D10年化收益", "年化波动", "最大回撤", "Sharpe", "年化超额", "月均换手"], top_rows, [1.75, 0.85, 0.85, 0.85, 0.65, 0.8, 0.75], numeric_cols={1, 2, 3, 4, 5, 6})

    add_heading(doc, "5.2 多空与分层单调性", 2)
    ls_rows = [
        [row["composite_factor"], pct(row["annual_return"]), pct(row["annual_volatility"]), num(row["sharpe"]), pct(row["max_drawdown"]), pct(row["win_rate"]), pct(row["cumulative_return"])]
        for _, row in long_short.sort_values("annual_return", ascending=False).iterrows()
    ]
    add_table(doc, ["组合因子", "D10-D1年化", "年化波动", "Sharpe", "最大回撤", "胜率", "累计收益"], ls_rows, [1.8, 0.8, 0.8, 0.65, 0.8, 0.75, 0.8], numeric_cols={1, 2, 3, 4, 5, 6})
    summary_rows = [
        [row["composite_factor"], pct(row["average_spearman_decile_ic"]), pct(row["monotonic_up_period_share"]), pct(row["d1_mean_return"]), pct(row["d10_mean_return"])]
        for _, row in layer_summary.iterrows()
    ]
    add_table(doc, ["组合因子", "平均分层Spearman", "完全单调月份占比", "D1平均收益", "D10平均收益"], summary_rows, [1.8, 1.1, 1.2, 0.95, 0.95], numeric_cols={1, 2, 3, 4})
    add_para(doc, "RankIC 加权多因子的 D10-D1 年化收益为 4.82%，显著优于等权多因子的 -1.70%。但分层并非严格单调，说明综合得分仍存在噪声，需要进一步筛选因子池。")

    add_heading(doc, "6. 与单因子表现对比", 1)
    single_top_rows = []
    for _, row in single_top.sort_values("annual_return", ascending=False).head(8).iterrows():
        single_top_rows.append([row["factor"], pct(row["annual_return"]), pct(row["annual_volatility"]), pct(row["max_drawdown"]), num(row["sharpe"]), pct(row["average_monthly_turnover"])])
    for _, row in top.iterrows():
        single_top_rows.append([f"{row['composite_factor']} D10", pct(row["annual_return"]), pct(row["annual_volatility"]), pct(row["max_drawdown"]), num(row["sharpe"]), pct(row["average_monthly_turnover"])])
    add_table(doc, ["策略", "年化收益", "年化波动", "最大回撤", "Sharpe", "月均换手"], single_top_rows, [1.9, 0.9, 0.9, 0.9, 0.7, 0.9], numeric_cols={1, 2, 3, 4, 5})

    single_ls_rows = []
    for _, row in single_ls.sort_values("annual_return", ascending=False).head(8).iterrows():
        single_ls_rows.append([row["factor"], pct(row["annual_return"]), pct(row["annual_volatility"]), num(row["sharpe"]), pct(row["max_drawdown"]), pct(row["win_rate"])])
    for _, row in long_short.sort_values("annual_return", ascending=False).iterrows():
        single_ls_rows.append([row["composite_factor"], pct(row["annual_return"]), pct(row["annual_volatility"]), num(row["sharpe"]), pct(row["max_drawdown"]), pct(row["win_rate"])])
    add_table(doc, ["策略", "D10-D1年化", "年化波动", "Sharpe", "最大回撤", "胜率"], single_ls_rows, [1.9, 0.9, 0.9, 0.7, 0.9, 0.8], numeric_cols={1, 2, 3, 4, 5})
    add_para(doc, "从当前样本看，PE_TTM、Volatility 和 Dividend_Yield 的单因子表现强于多因子。这主要反映样本期内价值、低波和红利风格较强，并不必然说明多因子框架无效。")

    add_heading(doc, "7. 图表说明", 1)
    add_heading(doc, "7.1 单因子与多因子对比", 2)
    add_figure(doc, IMAGES / "single_vs_multifactor.png", "图 1：单因子与多因子年化收益对比", "左图比较单因子中性化 Top20% 与多因子 D10 的年化收益；右图比较单因子 D10-D1 与多因子 D10-D1 的年化收益。该图显示，多因子优于等权合成，但低于当前样本期最强单因子。")

    add_heading(doc, "7.2 多因子十分组收益", 2)
    add_figure(doc, IMAGES / "layer_period_returns.png", "图 2：多因子分层持有期收益", "该图用于检查综合得分与未来收益之间的分层关系。RankIC 加权版本整体优于等权版本，但不是每一层都严格递增，说明信号仍有改进空间。")

    add_heading(doc, "7.3 多空净值曲线", 2)
    add_figure(doc, IMAGES / "long_short_nav.png", "图 3：D10-D1 多空净值曲线", "多空净值曲线显示 RankIC 加权多因子长期优于等权合成。等权合成的多空净值表现较弱，说明简单平均不能有效筛选因子。")

    add_heading(doc, "7.4 RankIC 权重变化", 2)
    add_figure(doc, IMAGES / "rankic_weights.png", "图 4：滚动 RankIC 权重变化", "权重主要集中在低波动、股息率和估值因子上，和单因子表现结果一致。该图也显示了动态加权机制的风格轮动特征。")

    add_heading(doc, "8. 金融解释与风险提示", 1)
    add_para(doc, "多因子没有跑赢样本期最强单因子是合理的。单因子冠军往往受益于特定市场风格，而多因子合成会分散风格暴露。当前样本中价值、红利和低波动表现突出，因此 PE_TTM、Dividend_Yield 和 Volatility 的集中暴露带来更高收益。")
    add_bullets(
        doc,
        [
            "多因子的主要价值是降低单一因子失效风险，而不是保证样本内收益最高。",
            "RankIC 加权使用历史信息，因此不会提前知道未来最强因子，存在自然滞后。",
            "当前全量 11 因子合成仍包含弱因子，后续应构建精选因子池。",
            "当前样本只有 54 个调仓期，需继续进行更长历史和滚动样本外检验。",
        ],
    )

    add_heading(doc, "9. 后续研究计划", 1)
    add_bullets(
        doc,
        [
            "基于长期 RankIC 和经济含义筛选因子池，优先测试 PE_TTM、Dividend_Yield、Volatility、Revenue_Growth 等因子组合。",
            "补充分年度、分市场状态和分股票池结果，检验多因子在不同环境下的稳健性。",
            "设计滚动样本外冠军单因子对比，避免用全样本最强单因子形成事后选择偏差。",
            "引入更真实的交易约束，包括冲击成本、涨跌停限制、流动性容量和行业偏离约束。",
        ],
    )

    add_heading(doc, "10. 阶段性结论", 1)
    add_para(doc, "本阶段完成了从单因子检测到多因子合成与分层回测的扩展。结果表明，简单等权合成不可取，滚动 RankIC 加权能有效改善多因子的排序能力和组合表现。当前多因子未超过最强单因子，但该现象与样本期风格特征一致，金融逻辑上可以解释。后续应围绕精选因子池和样本外稳健性继续推进。")

    doc.core_properties.title = "A股多因子合成与分层回测研究"
    doc.core_properties.subject = "多因子研究报告"
    doc.core_properties.author = "share_quant research"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(add_docx_report())

