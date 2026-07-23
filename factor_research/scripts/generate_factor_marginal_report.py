"""Generate the Word report for the frozen factor marginal analysis.

The report is intentionally generated from the completed analysis artifacts.  It
does not open the research database or recompute factor signals, portfolios, or
weights.
"""

from __future__ import annotations

import argparse
import math
import shutil
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK = "1E1E1E"
GREY = "5A5A5A"
LIGHT_GREY = "F2F4F7"
LIGHT_BLUE = "DDEBF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_RED = "FCE4D6"

STRATEGY_LABELS = {
    "Full11_Equal": "原始11因子等权",
    "Fixed4_Equal": "固定4因子等权",
    "Grouped8_Equal": "经济属性分组",
}

STABILITY_LABELS = {
    "stable_positive": "稳定正向",
    "stable_negative": "稳定负向",
    "mixed": "方向混合",
}


def parse_args() -> argparse.Namespace:
    project = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--reference",
        type=Path,
        default=project / "outputs" / "multifactor" / "multifactor_research_report.docx",
    )
    parser.add_argument(
        "--data-dir",
        type=Path,
        default=project / "outputs" / "factor_marginal_analysis",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=project
        / "outputs"
        / "factor_marginal_analysis"
        / "factor_marginal_analysis_research_report.docx",
    )
    return parser.parse_args()


def fmt_pct(value: float, digits: int = 2, signed: bool = False) -> str:
    if pd.isna(value):
        return "—"
    if signed:
        return f"{value:+.{digits}%}"
    return f"{value:.{digits}%}"


def fmt_num(value: float, digits: int = 4, signed: bool = False) -> str:
    if pd.isna(value):
        return "—"
    if signed:
        return f"{value:+.{digits}f}"
    return f"{value:.{digits}f}"


def compound(series: pd.Series) -> float:
    return float((1.0 + series.fillna(0.0)).prod() - 1.0)


def paired_stats(series: pd.Series) -> dict[str, float]:
    values = series.dropna().astype(float)
    n = int(values.size)
    std = float(values.std(ddof=1)) if n > 1 else math.nan
    t_stat = float(values.mean() / (std / math.sqrt(n))) if n > 1 and std > 0 else math.nan
    return {
        "n": n,
        "mean": float(values.mean()),
        "median": float(values.median()),
        "win_rate": float((values > 0).mean()),
        "std": std,
        "t_stat": t_stat,
    }


def load_results(data_dir: Path) -> dict[str, pd.DataFrame]:
    names = [
        "leave_one_out_summary",
        "add_one_back_summary",
        "period_marginal_contribution",
        "factor_group_definition",
        "group_correlation_matrix",
        "grouped_composite_metrics",
        "grouped_layer_returns",
        "grouped_top20_returns",
        "coverage_diagnostics",
    ]
    return {name: pd.read_csv(data_dir / f"{name}.csv") for name in names}


def oos_layer_profiles(layer: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    data = layer.copy()
    data["attribution_date"] = data["attribution_date"].astype(int)
    data = data[data["attribution_date"] >= 20250101]
    profiles = (
        data.groupby(["factor", "decile"], as_index=False)["average_forward_return"]
        .mean()
        .sort_values(["factor", "decile"])
    )
    rows: list[dict[str, float | str]] = []
    for factor, group in profiles.groupby("factor"):
        group = group.sort_values("decile")
        d1 = float(group.loc[group["decile"] == 1, "average_forward_return"].iloc[0])
        d10 = float(group.loc[group["decile"] == 10, "average_forward_return"].iloc[0])
        monotonic = float(group["decile"].corr(group["average_forward_return"], method="spearman"))
        rows.append(
            {
                "factor": factor,
                "d1": d1,
                "d10": d10,
                "spread": d10 - d1,
                "monotonic_spearman": monotonic,
            }
        )
    return profiles, pd.DataFrame(rows)


def monthly_advantage(
    layer: pd.DataFrame, top20: pd.DataFrame
) -> tuple[pd.DataFrame, dict[str, dict[str, float]]]:
    layer = layer.copy()
    layer["attribution_date"] = layer["attribution_date"].astype(int)
    layer = layer[layer["attribution_date"] >= 20250101]
    extremes = layer[layer["decile"].isin([1, 10])]
    spread = (
        extremes.pivot_table(
            index=["factor", "attribution_month"],
            columns="decile",
            values="average_forward_return",
            aggfunc="mean",
        )
        .reset_index()
        .assign(d10_d1=lambda x: x[10] - x[1])
    )
    spread_wide = spread.pivot(
        index="attribution_month", columns="factor", values="d10_d1"
    )
    d10_delta = spread_wide["Grouped8_Equal"] - spread_wide["Full11_Equal"]

    daily = top20.copy()
    daily["trade_date"] = pd.to_datetime(daily["trade_date"].astype(str), format="%Y%m%d")
    daily = daily[daily["trade_date"] >= pd.Timestamp("2025-01-01")]
    daily["month"] = daily["trade_date"].dt.to_period("M").astype(str)
    monthly_top = (
        daily.groupby(["factor", "month"], as_index=False)["portfolio_return"]
        .agg(compound)
        .pivot(index="month", columns="factor", values="portfolio_return")
    )
    top_delta = monthly_top["Grouped8_Equal"] - monthly_top["Full11_Equal"]

    monthly = pd.concat(
        [
            d10_delta.rename("d10_d1_delta"),
            top_delta.rename("top20_net_return_delta"),
        ],
        axis=1,
        join="inner",
    ).reset_index(names="month")
    stats = {
        "D10-D1": paired_stats(monthly["d10_d1_delta"]),
        "Top20净收益": paired_stats(monthly["top20_net_return_delta"]),
    }
    return monthly, stats


def configure_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["figure.dpi"] = 140
    plt.rcParams["savefig.dpi"] = 180


def make_report_charts(
    data_dir: Path,
    profiles: pd.DataFrame,
    monthly: pd.DataFrame,
) -> tuple[Path, Path]:
    configure_matplotlib()
    image_dir = data_dir / "images"
    image_dir.mkdir(exist_ok=True)

    profile_path = image_dir / "report_oos_decile_profile.png"
    fig, ax = plt.subplots(figsize=(10.4, 5.4))
    palette = {
        "Full11_Equal": "#2E74B5",
        "Fixed4_Equal": "#C55A11",
        "Grouped8_Equal": "#70AD47",
    }
    for factor, group in profiles.groupby("factor"):
        group = group.sort_values("decile")
        ax.plot(
            group["decile"],
            group["average_forward_return"] * 100,
            marker="o",
            linewidth=2.0,
            label=STRATEGY_LABELS[factor],
            color=palette[factor],
        )
    ax.axhline(0, color="#999999", linewidth=0.8)
    ax.set_xticks(range(1, 11))
    ax.set_xlabel("收益分组（D1 低分，D10 高分）")
    ax.set_ylabel("样本外平均持有期收益（%）")
    ax.set_title("样本外十分组收益曲线")
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False, ncol=3, loc="upper left")
    fig.tight_layout()
    fig.savefig(profile_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    monthly_path = image_dir / "report_monthly_grouped_advantage.png"
    fig, axes = plt.subplots(2, 1, figsize=(10.4, 6.8), sharex=True)
    x = np.arange(len(monthly))
    colors_d10 = np.where(monthly["d10_d1_delta"] >= 0, "#70AD47", "#C55A11")
    colors_top = np.where(monthly["top20_net_return_delta"] >= 0, "#70AD47", "#C55A11")
    axes[0].bar(x, monthly["d10_d1_delta"] * 100, color=colors_d10, width=0.72)
    axes[0].axhline(0, color="#555555", linewidth=0.8)
    axes[0].set_ylabel("D10-D1差（百分点）")
    axes[0].set_title("经济属性分组相对原始11因子：逐月收益差")
    axes[0].grid(axis="y", alpha=0.2)
    axes[1].bar(x, monthly["top20_net_return_delta"] * 100, color=colors_top, width=0.72)
    axes[1].axhline(0, color="#555555", linewidth=0.8)
    axes[1].set_ylabel("Top20净收益差（百分点）")
    axes[1].grid(axis="y", alpha=0.2)
    axes[1].set_xticks(x)
    axes[1].set_xticklabels(monthly["month"], rotation=45, ha="right", fontsize=8)
    fig.tight_layout()
    fig.savefig(monthly_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return profile_path, monthly_path


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    font_size: float = 8.5,
    color: str = DARK,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths: list[float] | None = None,
    font_size: float = 8.3,
    first_col_left: bool = True,
    highlight_rows: dict[int, str] | None = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GREY)
        set_cell_text(
            cell,
            text,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT
            if first_col_left and idx == 0
            else WD_ALIGN_PARAGRAPH.CENTER,
            font_size=font_size + 0.3,
        )
        if widths:
            cell.width = Inches(widths[idx])
    repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if highlight_rows and row_index in highlight_rows:
            for cell in row.cells:
                set_cell_shading(cell, highlight_rows[row_index])
        for idx, text in enumerate(values):
            set_cell_text(
                row.cells[idx],
                text,
                align=WD_ALIGN_PARAGRAPH.LEFT
                if first_col_left and idx == 0
                else WD_ALIGN_PARAGRAPH.CENTER,
                font_size=font_size,
            )
            if widths:
                row.cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.35
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        paragraph.add_run(text[len(bold_lead) :])
    else:
        paragraph.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.add_run("• ").bold = True
    paragraph.add_run(text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.35) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(2)
    run = cap.add_run(caption)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(GREY)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("share_quant research  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for style_name, size, color in [
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12, DARK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    for section in doc.sections:
        section.header.paragraphs[0].text = "多因子边际贡献与经济属性分组研究报告"
        header_run = section.header.paragraphs[0].runs[0]
        header_run.font.size = Pt(8.5)
        header_run.font.color.rgb = RGBColor.from_string(GREY)
        footer = section.footer.paragraphs[0]
        footer.text = ""
        add_page_number(footer)


def build_report(
    reference: Path,
    data_dir: Path,
    output: Path,
    frames: dict[str, pd.DataFrame],
    profiles: pd.DataFrame,
    profile_stats: pd.DataFrame,
    monthly: pd.DataFrame,
    monthly_stats: dict[str, dict[str, float]],
    profile_chart: Path,
    monthly_chart: Path,
) -> None:
    doc = Document(reference)
    clear_document_body(doc)
    configure_document(doc)

    groups = frames["factor_group_definition"]
    metrics = frames["grouped_composite_metrics"]
    loo = frames["leave_one_out_summary"]
    addback = frames["add_one_back_summary"]
    periods = frames["period_marginal_contribution"]
    corr = frames["group_correlation_matrix"].set_index("group_id")
    oos = metrics[metrics["period"] == "OOS_FULL"].set_index("strategy_id")

    full = oos.loc["Full11_Equal"]
    fixed = oos.loc["Fixed4_Equal"]
    grouped = oos.loc["Grouped8_Equal"]
    mean_abs_corr = float(
        corr.where(~np.eye(len(corr), dtype=bool)).abs().stack().mean()
    )

    # Cover
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("研究报告")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("A股多因子边际贡献与经济属性分组验证")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DARK)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(
        "基于 Leave-one-out、Add-one-back 与预冻结经济属性分组的样本外研究"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(GREY)

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["研究范围", "完整11因子、固定4因子、11个删除实验、7个加入实验及8组经济属性组合"],
            ["训练期", "2022-01-01 至 2024-12-31"],
            ["样本外期", "2025-01-01 至 2026-07-07（19个调仓月）"],
            ["数据与交易口径", "月末信号、下一交易日执行；数据库只读；Top20扣单边0.10%换手成本"],
            ["报告日期", date.today().isoformat()],
        ],
        widths=[1.35, 5.15],
        font_size=9.0,
    )

    add_heading(doc, "摘要")
    add_body(
        doc,
        "本研究针对固定4因子筛选方案在2025年后的样本外失效，重新从组合边际贡献和经济属性分散角度解释原始11因子等权组合的优势。研究严格冻结因子方向、分组、权重、信号时点与交易口径，不使用样本外收益重新选择因子或调参。"
    )
    add_body(
        doc,
        f"样本外结果显示，原始11因子的D10-D1年化为{fmt_pct(full['d10_d1_annual_return'])}，"
        f"Top20年化超额为{fmt_pct(full['top20_excess_annual_return'])}；固定4因子对应为"
        f"{fmt_pct(fixed['d10_d1_annual_return'])}和{fmt_pct(fixed['top20_excess_annual_return'])}。"
        "固定4因子的平均RankIC更高，但覆盖率更低且月度波动更大，说明单一排序统计量不足以代表组合层面的稳健贡献。"
    )
    add_body(
        doc,
        f"预冻结的经济属性分组组合保留全部11因子，采用“组内等权、组间等权”。其样本外D10-D1年化为"
        f"{fmt_pct(grouped['d10_d1_annual_return'])}，Top20年化超额为"
        f"{fmt_pct(grouped['top20_excess_annual_return'])}，分别高于原始11因子"
        f"{fmt_pct(grouped['d10_d1_annual_return']-full['d10_d1_annual_return'], signed=True)}和"
        f"{fmt_pct(grouped['top20_excess_annual_return']-full['top20_excess_annual_return'], signed=True)}。"
        "然而，逐月配对优势的简单t统计量均低于1，且分组组合换手更高、Top20最大回撤略差，因此结果支持其成为前瞻验证的挑战者，而不足以立即替换原始11因子等权基准。"
    )
    add_bullet(doc, "组合价值最明确的单因子是 Revenue_Growth：在Leave-one-out的两个主要指标上均呈稳定正向贡献。")
    add_bullet(doc, "原始11因子优于固定4因子，主要来自被筛除因子的互补排序、经济风格分散和更完整的截面覆盖。")
    add_bullet(doc, "经济属性分组在全样本外同时改善两个主要收益指标，但并未在风险、换手和统计显著性上形成全面支配。")
    add_bullet(doc, "正式决策建议：保留原始11因子等权为生产基准，将分组方案冻结为 GroupBalanced11_v1，继续前瞻观察。")

    # 1
    add_heading(doc, "1. 研究背景与目标")
    add_body(
        doc,
        "此前研究依据2022—2024年训练期的经济方向、RankIC均值、胜率、稳定性、覆盖率和相关性，从11个原始因子中筛选出PE_TTM、Dividend_Yield、Volatility与Main_Net_In。该固定4因子组合在2025年后的D10-D1与Top20超额收益均明显弱于原始11因子等权，因而筛选方案已经判定失败。"
    )
    add_body(doc, "本阶段不再寻找样本外最优子集，而是回答以下五个预定义问题：")
    add_bullet(doc, "各因子相对不同组合基准的边际贡献方向是什么。")
    add_bullet(doc, "原始11因子为何能够明显优于固定4因子。")
    add_bullet(doc, "按经济属性先组内等权、再组间等权，能否同时改善D10-D1与Top20超额收益。")
    add_bullet(doc, "上述改善是否跨2025H1、2025H2和2026YTD保持方向稳定。")
    add_bullet(doc, "现有证据是否足以替换原始11因子等权组合。")

    # 2
    add_heading(doc, "2. 研究设计与回测口径")
    add_table(
        doc,
        ["项目", "冻结设置"],
        [
            ["因子与方向", "来自既有walk_forward_champion配置的完整11因子，不重新定义"],
            ["训练/样本外隔离", "训练期2022—2024；主要评价期自2025-01-01起；样本外收益不进入权重"],
            ["信号与执行", "月末生成信号，下一交易日执行；沿用中性化、停牌和缺失收益处理"],
            ["D10-D1", "D10平均远期毛收益减D1平均远期毛收益，月度复利年化；不扣成本"],
            ["Top20超额", "最高20%股票等权；扣单边0.10%换手成本；相对000985.CSI年化超额"],
            ["辅助指标", "RankIC、最大回撤、换手率、覆盖率和预定义子区间稳定性"],
            ["边际贡献", "Leave-one-out：Full11－删除方案；Add-one-back：加入方案－Fixed4"],
            ["稳定标签", "全样本外同号，且三个子区间中至少两个同号"],
            ["数据库边界", "研究脚本以read_only=True连接；本报告仅读取既有CSV和PNG，不连接数据库"],
        ],
        widths=[1.35, 5.15],
        font_size=8.7,
    )
    add_body(
        doc,
        "D10-D1与Top20超额收益是并列的主要指标，不合并为单一评分。最大回撤的“改善”应理解为回撤绝对值缩小；边际贡献属于特定基准附近的局部差异，不具有Shapley可加性，也不能被直接解释为因果效应。"
    )

    # 3
    add_heading(doc, "3. 因子池与经济属性分组")
    group_rows = []
    for _, row in groups.iterrows():
        direction = "正向" if int(row["direction"]) == 1 else "反向"
        group_rows.append(
            [
                row["factor"],
                direction,
                row["group_name_zh"],
                fmt_pct(1 / 11),
                fmt_pct(row["effective_factor_weight_full_data"]),
                "是" if bool(row["fixed_pool_member"]) else "否",
            ]
        )
    add_table(
        doc,
        ["因子", "方向", "经济组", "11因子权重", "分组有效权重", "固定4"],
        group_rows,
        widths=[1.45, 0.65, 0.85, 1.05, 1.15, 0.7],
        font_size=7.8,
    )
    add_body(
        doc,
        "八个预冻结经济组分别为价值、质量、成长、低波、动量、资金流、流动性和股东结构。没有与现有因子直接对应的规模组和纯情绪组，故未强行归类。分组组合先在组内对可用因子等权，再在有效组之间等权；完整数据下，每个组权重为12.5%。"
    )
    add_body(
        doc,
        "该规则改变的是经济属性的预算，而不是因子名单：Revenue_Growth、Volatility、Momentum_60D、Main_Net_In、Turnover_20D和Holder_Concen等单因子组各获得12.5%；价值组两个因子各6.25%；质量组三个因子各4.17%。"
    )

    # 4
    add_heading(doc, "4. 样本外核心结果")
    core_rows = []
    for strategy_id in ["Full11_Equal", "Fixed4_Equal", "Grouped8_Equal"]:
        row = oos.loc[strategy_id]
        core_rows.append(
            [
                STRATEGY_LABELS[strategy_id],
                fmt_pct(row["d10_d1_annual_return"]),
                fmt_pct(row["top20_excess_annual_return"]),
                fmt_num(row["rank_ic_mean"]),
                fmt_pct(row["d10_d1_max_drawdown"]),
                fmt_pct(row["top20_max_drawdown"]),
                fmt_pct(row["d10_d1_average_two_leg_turnover"]),
                fmt_pct(row["top20_average_monthly_turnover"]),
                fmt_pct(row["average_valid_stock_coverage"]),
            ]
        )
    add_table(
        doc,
        ["组合", "D10-D1", "Top20超额", "RankIC", "D10回撤", "Top20回撤", "双腿换手", "Top20换手", "覆盖率"],
        core_rows,
        widths=[1.1, 0.68, 0.72, 0.57, 0.68, 0.72, 0.68, 0.68, 0.58],
        font_size=7.2,
        highlight_rows={0: LIGHT_BLUE, 2: LIGHT_GREEN},
    )
    add_picture(
        doc,
        data_dir / "images" / "grouped_composite_comparison.png",
        "图 1：原始11因子、固定4因子与经济属性分组组合的样本外指标对比",
    )
    add_body(
        doc,
        "固定4因子的平均RankIC为4.81%，高于原始11因子的0.94%，但D10-D1和Top20超额均为负。这并非口径错误：固定4因子的IC在少数月份很高、在弱势月份尾部更差；同时其平均覆盖率仅79.61%，明显低于原始11因子的99.90%。因此，平均RankIC不能替代组合收益、尾部稳定性和覆盖率评价。"
    )
    add_body(
        doc,
        "固定4因子的Top20最大回撤为-14.69%，好于原始11因子的-17.21%，体现价值、红利、低波和资金流组合具有一定防御属性。原始11因子不是所有风险指标都占优，其优势集中在两个主要收益指标与D10-D1尾部控制。"
    )

    profile_rows = []
    for strategy_id in ["Full11_Equal", "Fixed4_Equal", "Grouped8_Equal"]:
        row = profile_stats.set_index("factor").loc[strategy_id]
        profile_rows.append(
            [
                STRATEGY_LABELS[strategy_id],
                fmt_pct(row["d1"]),
                fmt_pct(row["d10"]),
                fmt_pct(row["spread"], signed=True),
                fmt_num(row["monotonic_spearman"], 3),
            ]
        )
    add_table(
        doc,
        ["组合", "D1平均收益", "D10平均收益", "平均月度D10-D1", "分层Spearman"],
        profile_rows,
        widths=[1.55, 1.15, 1.15, 1.35, 1.15],
        font_size=8.4,
    )
    add_picture(doc, profile_chart, "图 2：样本外十分组平均持有期收益曲线")
    add_body(
        doc,
        "原始11因子和经济属性分组的分层Spearman分别约为0.72和0.85，D10收益均明显高于D1；固定4因子约为-0.48，呈反向分层。该截面形态与两个主要收益指标一致，为“分散组合改善排序”的金融解释提供了直观证据。"
    )

    # 5
    add_heading(doc, "5. Leave-one-out：原始11因子的局部边际贡献")
    add_body(
        doc,
        "Leave-one-out将原始11因子等权作为基准，每次删除一个因子。表中正值表示“保留该因子”优于“删除该因子”；负值表示该因子在这一局部基准附近拖累对应指标。"
    )
    loo_rows = []
    for _, row in loo.sort_values("d10_d1_annual_return_delta", ascending=False).iterrows():
        loo_rows.append(
            [
                row["factor"],
                fmt_pct(row["d10_d1_annual_return_delta"], signed=True),
                STABILITY_LABELS[row["d10_d1_annual_return_stability"]],
                fmt_pct(row["top20_excess_annual_return_delta"], signed=True),
                STABILITY_LABELS[row["top20_excess_annual_return_stability"]],
                fmt_num(row["rank_ic_mean_delta"], signed=True),
                fmt_pct(row["top20_average_monthly_turnover_delta"], signed=True),
            ]
        )
    add_table(
        doc,
        ["因子", "D10-D1边际", "D10稳定性", "Top20边际", "Top20稳定性", "RankIC边际", "Top20换手边际"],
        loo_rows,
        widths=[1.32, 0.83, 0.85, 0.83, 0.85, 0.83, 0.93],
        font_size=7.3,
    )
    add_picture(
        doc,
        data_dir / "images" / "leave_one_out_primary_metrics.png",
        "图 3：Leave-one-out两个主要指标的边际贡献",
        width=6.45,
    )
    add_body(
        doc,
        "Revenue_Growth是唯一在Leave-one-out的D10-D1与Top20超额上均为稳定正向的因子，边际分别为+7.26和+3.52个百分点，可标记为“可能具有稳定组合价值”。Momentum_60D、PE_TTM、Holder_Concen和Turnover_20D对Top20有稳定正向贡献，但D10-D1结论混合或相反。"
    )
    add_body(
        doc,
        "Dividend_Yield与Volatility在两个主要指标上均呈稳定负向局部贡献；但这不等于应立即删除。边际贡献依赖其余因子的权重和基准构成，且二者可能提供防御、回撤或风格约束价值。基于本次样本外结果删除因子会构成事后选择。"
    )

    # 6
    add_heading(doc, "6. Add-one-back：固定4因子失效的来源")
    add_body(
        doc,
        "Add-one-back以固定4因子为基准，每次加入一个此前被筛除的因子。该实验回答“哪些缺失成分能够修复固定4因子”的问题，不用于事后构建最优5因子池。"
    )
    addback_rows = []
    for _, row in addback.sort_values("d10_d1_annual_return_delta", ascending=False).iterrows():
        addback_rows.append(
            [
                row["factor"],
                fmt_pct(row["d10_d1_annual_return_delta"], signed=True),
                STABILITY_LABELS[row["d10_d1_annual_return_stability"]],
                fmt_pct(row["top20_excess_annual_return_delta"], signed=True),
                STABILITY_LABELS[row["top20_excess_annual_return_stability"]],
                fmt_num(row["rank_ic_mean_delta"], signed=True),
                fmt_pct(row["average_valid_stock_count_delta"] / row["excluded_average_valid_stock_count"]),
            ]
        )
    add_table(
        doc,
        ["加入因子", "D10-D1边际", "D10稳定性", "Top20边际", "Top20稳定性", "RankIC边际", "覆盖股票增幅"],
        addback_rows,
        widths=[1.32, 0.83, 0.85, 0.83, 0.85, 0.83, 0.93],
        font_size=7.3,
    )
    add_picture(
        doc,
        data_dir / "images" / "add_one_back_primary_metrics.png",
        "图 4：Add-one-back两个主要指标的边际贡献",
        width=6.45,
    )
    add_body(
        doc,
        "七个被筛除因子加入固定4后，D10-D1均呈稳定正向；除Turnover_20D外，其余六个因子的Top20超额也稳定改善。Revenue_Growth、Momentum_60D和ROE对D10-D1的修复幅度最大，分别约为+18.45、+14.63和+10.22个百分点。"
    )
    add_body(
        doc,
        "该结果说明固定4因子的失败并非由某一个遗漏因子造成，而是过度压缩了经济属性和截面覆盖。被筛除因子即使单因子RankIC较弱，仍可能通过不同收益来源、风险状态或缺失结构提供组合互补性。"
    )

    # 7
    add_heading(doc, "7. 经济属性分散与训练期相关性")
    add_body(
        doc,
        f"训练期使用每月截面中性化组得分计算Spearman相关，再对月份取均值。八组之间的平均绝对非对角相关系数为{mean_abs_corr:.4f}，总体较低，说明分组并非仅在名称上分散。"
    )
    corr_pairs = []
    columns = list(corr.columns)
    for i, left in enumerate(columns):
        for right in columns[i + 1 :]:
            corr_pairs.append((left, right, float(corr.loc[left, right])))
    corr_pairs.sort(key=lambda x: abs(x[2]), reverse=True)
    add_table(
        doc,
        ["组别一", "组别二", "训练期相关系数", "解释"],
        [
            [
                left,
                right,
                fmt_num(value, 3, signed=True),
                "存在较强互补" if value < -0.25 else "存在一定同向暴露",
            ]
            for left, right, value in corr_pairs[:5]
        ],
        widths=[1.25, 1.25, 1.35, 2.65],
        font_size=8.5,
    )
    add_picture(
        doc,
        data_dir / "images" / "group_correlation_heatmap.png",
        "图 5：训练期经济属性组间月均截面Spearman相关性",
        width=5.9,
    )
    add_body(
        doc,
        "绝对值最大的关系是低波与流动性约-0.50，其次是动量与流动性约+0.28、价值与低波约+0.27。分组等权降低了价值组和质量组内部多因子对总组合的机械权重，同时提升成长、动量、资金流、流动性和股东结构等单因子组的风格预算。"
    )

    # 8
    add_heading(doc, "8. 子区间稳定性与逐月证据")
    group_period = periods[
        (periods["included_strategy"] == "Grouped8_Equal")
        & (periods["excluded_strategy"] == "Full11_Equal")
    ].copy()
    if group_period.empty:
        group_period = (
            metrics[metrics["period"].isin(["2025H1", "2025H2", "2026YTD"])]
            .pivot(index="period", columns="strategy_id")
            .reset_index()
        )
        period_rows = []
        for period in ["2025H1", "2025H2", "2026YTD"]:
            period_rows.append(
                [
                    period,
                    fmt_pct(
                        group_period.loc[
                            group_period["period"] == period,
                            ("d10_d1_annual_return", "Grouped8_Equal"),
                        ].iloc[0]
                        - group_period.loc[
                            group_period["period"] == period,
                            ("d10_d1_annual_return", "Full11_Equal"),
                        ].iloc[0],
                        signed=True,
                    ),
                    fmt_pct(
                        group_period.loc[
                            group_period["period"] == period,
                            ("top20_excess_annual_return", "Grouped8_Equal"),
                        ].iloc[0]
                        - group_period.loc[
                            group_period["period"] == period,
                            ("top20_excess_annual_return", "Full11_Equal"),
                        ].iloc[0],
                        signed=True,
                    ),
                ]
            )
    else:
        period_rows = [
            [
                row["period"],
                fmt_pct(row["d10_d1_annual_return_delta"], signed=True),
                fmt_pct(row["top20_excess_annual_return_delta"], signed=True),
            ]
            for _, row in group_period.iterrows()
            if row["period"] in ["2025H1", "2025H2", "2026YTD"]
        ]
    add_table(
        doc,
        ["样本外子区间", "分组－11因子 D10-D1", "分组－11因子 Top20超额"],
        period_rows,
        widths=[1.7, 2.4, 2.4],
        font_size=8.7,
    )
    add_picture(
        doc,
        data_dir / "images" / "subperiod_stability.png",
        "图 6：主要方案在预定义样本外子区间的表现",
    )
    add_body(
        doc,
        "分组组合相对原始11因子的D10-D1在三个子区间均为正；Top20超额在2025H2与2026YTD为正，在2025H1为负。因此，按预设稳定标签，D10-D1属于稳定改善，Top20属于“全样本外改善且2/3子区间同向”，也满足方向稳定规则，但并非每一段都占优。"
    )

    monthly_rows = []
    for metric, stats in monthly_stats.items():
        monthly_rows.append(
            [
                metric,
                str(int(stats["n"])),
                fmt_pct(stats["mean"], signed=True),
                fmt_pct(stats["median"], signed=True),
                fmt_pct(stats["win_rate"]),
                fmt_pct(stats["std"]),
                fmt_num(stats["t_stat"], 2),
            ]
        )
    add_table(
        doc,
        ["月度配对指标", "月份数", "均值差", "中位数差", "胜率", "月度标准差", "简单t值"],
        monthly_rows,
        widths=[1.25, 0.65, 0.9, 0.9, 0.75, 0.95, 0.8],
        font_size=8.2,
    )
    add_picture(
        doc,
        monthly_chart,
        "图 7：经济属性分组相对原始11因子的逐月收益差（绿色为占优）",
    )
    add_body(
        doc,
        "19个月配对结果中，分组方案的D10-D1月均优势约为+0.20个百分点，月度胜率52.63%，简单t值约0.94；Top20净收益月均优势约+0.09个百分点，胜率63.16%，简单t值约0.69。该统计量未进行HAC调整，且样本较短，不能据此声称收益差具有传统统计显著性。"
    )

    # 9
    add_heading(doc, "9. 金融逻辑、实施成本与风险提示")
    add_body(doc, "结果整体符合多因子投资的金融逻辑，主要体现在以下方面：")
    add_bullet(
        doc,
        "风格分散而非单因子冠军复制。原始11因子覆盖价值、质量、成长、低波、动量、资金行为、流动性和股东结构，固定4因子则集中在价值、红利、低波与资金流。"
    )
    add_bullet(
        doc,
        "弱RankIC不等于零组合价值。因子可以通过覆盖不同股票、缓冲特定市场状态或改善尾部月份，对组合产生互补贡献。"
    )
    add_bullet(
        doc,
        "高平均RankIC与负组合收益可以并存。固定4因子的RankIC波动更大，收益分层呈反向，且较低覆盖率改变了可交易横截面。"
    )
    add_bullet(
        doc,
        "防御性与收益性存在权衡。固定4因子的Top20最大回撤更小；分组组合虽然提高收益，却使D10双腿换手和Top20月均换手分别约提高27%和24%，Top20最大回撤也略深。"
    )
    add_body(
        doc,
        "经济属性分组的改善并不证明“八组权重”是结构性最优。它只证明，在本次预先冻结的规则和当前样本外窗口中，减少组内重复计票、提高独立经济属性预算具有合理效果。任何进一步删因子、调组权重或选择表现最佳子区间，都会重新引入样本外数据挖掘。"
    )

    # 10
    add_heading(doc, "10. 结论与组合决策")
    decision_rows = [
        [
            "1. 哪些因子具有正/负边际贡献",
            "Revenue_Growth在LOO两个主要指标均稳定正向；其余因子多表现为指标依赖或基准依赖。Dividend_Yield与Volatility在LOO两个主要指标均稳定负向，但不足以据此事后删除。",
        ],
        [
            "2. 原始11因子为何优于固定4因子",
            "被筛除因子提供互补排序、风格分散与更高覆盖率；固定4因子过度集中，RankIC虽高但分层、尾部月份和组合收益较弱。",
        ],
        [
            "3. 分组是否同时改善主要指标",
            f"是。全样本外D10-D1提高{fmt_pct(grouped['d10_d1_annual_return']-full['d10_d1_annual_return'], signed=True)}，Top20年化超额提高{fmt_pct(grouped['top20_excess_annual_return']-full['top20_excess_annual_return'], signed=True)}。",
        ],
        [
            "4. 是否跨子区间稳定",
            "D10-D1在三个子区间均改善；Top20在两个子区间改善、2025H1落后。按预设2/3规则稳定，但非全区间一致。",
        ],
        [
            "5. 是否有充分证据替换11因子",
            "尚无。方向规则支持分组方案成为挑战者，但19个月样本、低配对t值、更高换手和略差Top20回撤不足以支持生产替换。",
        ],
    ]
    add_table(
        doc,
        ["决策问题", "结论"],
        decision_rows,
        widths=[2.05, 4.45],
        font_size=8.7,
        highlight_rows={4: LIGHT_RED},
    )
    add_body(
        doc,
        "建议将原始11因子等权继续作为生产/研究基准，将本次预冻结分组方案命名并冻结为 GroupBalanced11_v1，保持全部11因子、组内等权和组间等权，不再依据当前样本外结果调整。后续只做前瞻验证，至少积累新的完整市场阶段，再依据预先写定的收益、回撤、换手和显著性门槛决定是否替换。"
    )
    add_body(
        doc,
        "当前最稳健的结论不是“已经找到必然更强的组合”，而是“找到一个金融逻辑清晰、两个主要指标均占优、但尚需更长样本确认的候选组合”。"
    )

    # Appendix
    add_heading(doc, "附录：可复核产物与审计")
    add_table(
        doc,
        ["产物", "用途"],
        [
            ["leave_one_out_summary.csv", "11个因子相对原始组合的全样本外边际贡献"],
            ["add_one_back_summary.csv", "7个被筛除因子相对固定4因子的修复贡献"],
            ["monthly_marginal_contribution.csv", "逐月边际贡献明细"],
            ["period_marginal_contribution.csv", "2025H1、2025H2、2026YTD分期结果"],
            ["factor_group_definition.csv", "预冻结经济属性分组与有效权重"],
            ["group_correlation_matrix.csv", "训练期组间月均截面相关性"],
            ["grouped_composite_metrics.csv", "三个核心组合的全期和分期指标"],
            ["grouped_layer_returns.csv", "十分组持有期收益明细"],
            ["grouped_top20_returns.csv", "Top20日频净收益与基准收益"],
            ["protected_file_audit.csv", "数据库、旧代码与旧输出的纳秒级时间戳/大小审计"],
        ],
        widths=[2.15, 4.35],
        font_size=8.4,
    )
    add_body(
        doc,
        "完整运行共生成21个策略（11个Leave-one-out、7个Add-one-back及3个核心组合）。数据库连接模式为read_only；受保护文件审计覆盖17,169个文件且全部未变。测试与编译验证已通过。"
    )
    add_body(
        doc,
        "本报告由既有分析CSV和PNG生成，不重新访问数据库，也不改变任何策略规则。报告中的t值为描述性简单配对统计量，未用于事后筛选或权重优化。"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    temp = output.with_suffix(".tmp.docx")
    doc.save(temp)
    shutil.move(temp, output)


def main() -> None:
    args = parse_args()
    frames = load_results(args.data_dir)
    profiles, profile_stats = oos_layer_profiles(frames["grouped_layer_returns"])
    monthly, monthly_stats = monthly_advantage(
        frames["grouped_layer_returns"], frames["grouped_top20_returns"]
    )
    profile_chart, monthly_chart = make_report_charts(
        args.data_dir, profiles, monthly
    )
    build_report(
        args.reference,
        args.data_dir,
        args.output,
        frames,
        profiles,
        profile_stats,
        monthly,
        monthly_stats,
        profile_chart,
        monthly_chart,
    )
    print(args.output)


if __name__ == "__main__":
    main()
