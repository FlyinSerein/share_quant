from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = Path(__file__).resolve().parents[1] / "outputs" / "multifactor" / "multifactor_research_report_for_mentor.docx"
doc = Document(path)
text = "\n".join(p.text for p in doc.paragraphs)
checks = [
    "A股多因子合成与分层回测研究",
    "研究背景与目标",
    "多因子构造方法",
    "图表说明",
    "阶段性结论",
]
print("paragraphs", len(doc.paragraphs))
print("tables", len(doc.tables))
print("inline_shapes", len(doc.inline_shapes))
print("sections", len(doc.sections))
for item in checks:
    print(item, item in text)
with ZipFile(path) as z:
    media = [name for name in z.namelist() if name.startswith("word/media/")]
    print("media_count", len(media))
    print("document_xml", "word/document.xml" in z.namelist())

for forbidden in ["一句话结论", "多因子的核心价值", "给导师的阶段性总结"]:
    print("forbidden", forbidden, forbidden in text)
