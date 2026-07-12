from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "outputs" / "multifactor"
TABLES = BASE / "tables"
IMAGES = BASE / "images"
OUT = BASE / "multifactor_research_report_for_mentor.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
LIGHT_BLUE_FILL = "E8EEF5"
RISK_FILL = "FFF4E5"


def pct(value: object) -> str:
    if pd.isna(value):
        return ""
    return f"{float(value):.2%}"


def num(value: object, digits: int = 2) -> str:
    if pd.isna(value):
        return ""
    return f"{float(value):.{digits}f}"


def set_run_font(run, name: str = "Calibri", east_asia: str = "Microsoft YaHei", size: float | None = None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para(p, before: float = 0, after: float = 6, line_spacing: float = 1.10, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align


def add_para(doc, text: str = "", size: float = 11, color=INK, bold: bool = False, italic: bool = False, before: float = 0, after: float = 6, align=None):
    p = doc.add_paragraph()
    set_para(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level: int):
    style_name = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style_name)
    if level == 1:
        set_para(p, before=16, after=8)
        size, color = 16, BLUE
    elif level == 2:
        set_para(p, before=12, after=6)
        size, color = 13, BLUE
    else:
        set_para(p, before=8, after=4)
        size, color = 12, DARK_BLUE
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_table(table, widths_in: list[float], numeric_cols: set[int] | None = None):
    numeric_cols = numeric_cols or set()
    set_table_geometry(table, widths_in)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                set_para(p, after=2, line_spacing=1.10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=9.2 if i else 9.5, color=INK, bold=(i == 0))
            if i == 0:
                shade_cell(cell, HEADER_FILL)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_in: list[float], numeric_cols: set[int] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
    style_table(table, widths_in, numeric_cols)
    add_para(doc, "", after=2)
    return table


def add_callout(doc, title: str, body: str, fill: str = CALLOUT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    set_para(p, after=4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_para(p2, after=0, line_spacing=1.15)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    set_table_geometry(table, [6.5])
    add_para(doc, "", after=2)


def add_figure(doc, image_path: Path, caption: str, note: str, width: float = 6.25):
    doc.add_picture(str(image_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, caption, size=9.5, color=MUTED, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, note, size=10.2, color=INK, after=8)


def add_bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style=None)
        set_para(p, after=4, line_spacing=1.167)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        r = p.add_run("• ")
        set_run_font(r, size=10.5, color=BLUE, bold=True)
        r2 = p.add_run(item)
        set_run_font(r2, size=10.5, color=INK)


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "多因子合成研究报告"
    set_para(header, after=0)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "share_quant research | generated for internship mentor review"
    set_para(footer, after=0)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc: Document):
    add_para(doc, "研究报告", size=11, color=BLUE, bold=True, after=4)
    add_para(doc, "A股多因子合成与分层回测研究", size=24, color=INK, bold=True, after=6)
    add_para(doc, "基于单因子中性化得分、滚动 RankIC 加权与月度分层回测", size=13, color=MUTED, after=12)

    rows = [
        ["报告对象", "实习导师"],
        ["研究范围", "11 个 A 股单因子、多因子等权合成、滚动 RankIC 加权合成"],
        ["样本区间", "2022-01-01 至 2026-07-07"],
        ["数据与代码", "本地 share_quant 研究数据库；输出目录 research/factor_combo/outputs/multifactor"],
        ["生成日期", str(date.today())],
    ]
    add_table(doc, ["项目", "说明"], rows, [1.4, 5.1])


def build_report():
    single_top = pd.read_csv(TABLES / "single_factor_top_metrics.csv")
    single_ls = pd.read_csv(TABLES / "single_factor_long_short_metrics.csv")
    layer_metrics = pd.read_csv(TABLES / "layer_metrics.csv")
    long_short = pd.read_csv(TABLES / "long_short_metrics.csv")
    coverage = pd.read_csv(TABLES / "composite_scores_coverage.csv")
    ic = pd.read_csv(TABLES / "composite_ic_summary.csv")
    weights = pd.read_csv(TABLES / "composite_factor_weights.csv")

    doc = Document()
    setup_document(doc)
    add_cover(doc)

    add_heading(doc, "1. 研究目的与问题定义", 1)
    add_para(
        doc,
        "本研究在已有单因子获取、检测和回测框架基础上，进一步检验多因子合成是否能提升组合稳健性。核心问题不是单纯比较某一期样本内收益高低，而是判断：多因子是否能够降低弱因子拖累、改善排序能力，并在金融逻辑上解释其与强单因子的差异。",
    )
    add_bullets(
        doc,
        [
            "输入信号为 11 个单因子的行业和市值中性化得分。",
            "合成方式包括等权合成和滚动 RankIC 加权合成。",
            "评价指标包括 Top 层多头收益、D10-D1 多空收益、RankIC、换手、回撤和单因子对比。",
        ]
    )

    add_heading(doc, "2. 方法与回测口径", 1)
    method_rows = [
        ["股票池", "已上市、非停牌、非 ST；每期至少 6 个有效单因子得分"],
        ["信号处理", "方向统一、去极值、标准化；随后对行业和市值暴露做中性化"],
        ["调仓方式", "月末生成信号，下一交易日成交，持有至下一次调仓"],
        ["分层方式", "按得分分为 10 组，D10 为最高得分组；D10-D1 衡量横截面排序能力"],
        ["交易成本", "单边 10bp，在组合换手时扣除"],
        ["基准", "中证全指 000985.CSI"],
    ]
    add_table(doc, ["维度", "设定"], method_rows, [1.5, 5.0])

    add_heading(doc, "3. 样本覆盖与数据质量", 1)
    coverage_rows = []
    for _, row in coverage.iterrows():
        coverage_rows.append(
            [
                row["factor"],
                f"{int(row['signal_count'])}",
                str(int(row["first_signal"])),
                str(int(row["last_signal"])),
                f"{row['average_valid_stocks']:.0f}",
                f"{row['average_available_factor_count']:.2f}",
            ]
        )
    add_table(doc, ["组合因子", "期数", "首个信号", "最后信号", "平均股票数", "平均因子数"], coverage_rows, [1.7, 0.7, 1.0, 1.0, 1.0, 1.1], numeric_cols={1, 2, 3, 4, 5})
    add_para(doc, "覆盖结果显示，每期平均接近 5000 只股票，平均可用因子数超过 10 个，说明结果不是由少量样本或严重缺失驱动。", size=10.5)

    add_heading(doc, "4. 多因子核心结果", 1)
    top = layer_metrics[layer_metrics["bucket"].eq(10)].sort_values("annual_return", ascending=False)
    top_rows = [
        [
            row["composite_factor"],
            pct(row["annual_return"]),
            pct(row["annual_volatility"]),
            pct(row["max_drawdown"]),
            num(row["sharpe"]),
            pct(row["excess_annual_return"]),
            pct(row["average_monthly_turnover"]),
        ]
        for _, row in top.iterrows()
    ]
    add_table(doc, ["组合因子", "D10年化", "年化波动", "最大回撤", "Sharpe", "年化超额", "月均换手"], top_rows, [1.75, 0.8, 0.85, 0.85, 0.65, 0.85, 0.85], numeric_cols={1, 2, 3, 4, 5, 6})

    ls_rows = [
        [
            row["composite_factor"],
            pct(row["annual_return"]),
            pct(row["annual_volatility"]),
            num(row["sharpe"]),
            pct(row["max_drawdown"]),
            pct(row["win_rate"]),
        ]
        for _, row in long_short.sort_values("annual_return", ascending=False).iterrows()
    ]
    add_table(doc, ["组合因子", "D10-D1年化", "年化波动", "Sharpe", "最大回撤", "胜率"], ls_rows, [1.9, 0.9, 0.9, 0.7, 0.9, 0.8], numeric_cols={1, 2, 3, 4, 5})

    add_callout(
        doc,
        "结果解读",
        "滚动 RankIC 加权多因子在 D10 多头收益、D10-D1 多空收益、回撤和换手上均优于等权多因子，说明动态权重机制有效降低了弱因子拖累。",
    )

    add_heading(doc, "5. 与单因子的横向比较", 1)
    doc.add_picture(str(IMAGES / "single_vs_multifactor.png"), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图 1：单因子中性化 Top20% 与多因子 D10，以及单因子 D10-D1 与多因子 D10-D1 的年化收益对比。", size=9.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    single_top_rows = []
    for _, row in single_top.sort_values("annual_return", ascending=False).head(5).iterrows():
        single_top_rows.append([row["factor"], pct(row["annual_return"]), pct(row["annual_volatility"]), pct(row["max_drawdown"]), num(row["sharpe"]), pct(row["average_monthly_turnover"])])
    for _, row in top.iterrows():
        single_top_rows.append([f"{row['composite_factor']} D10", pct(row["annual_return"]), pct(row["annual_volatility"]), pct(row["max_drawdown"]), num(row["sharpe"]), pct(row["average_monthly_turnover"])])
    add_table(doc, ["策略", "年化收益", "年化波动", "最大回撤", "Sharpe", "月均换手"], single_top_rows, [1.9, 0.9, 0.9, 0.9, 0.7, 0.9], numeric_cols={1, 2, 3, 4, 5})

    add_para(
        doc,
        "从多头收益看，PE_TTM、Volatility、Dividend_Yield 等单因子在当前样本期强于多因子。该现象并不代表多因子失效，而是说明样本期价值、低波和红利风格占优，单因子集中暴露在这些风格上时收益更高。",
    )

    add_heading(doc, "6. 为什么多因子没有跑赢最强单因子", 1)
    add_bullets(
        doc,
        [
            "样本期存在明显风格偏向。PE_TTM、低波动、股息率直接受益于价值、低波和红利风格，因此单因子表现突出。",
            "多因子不是事后选择冠军因子，而是同时纳入多个因子。弱因子和阶段性失效因子即使权重较低，也会稀释最强因子的收益。",
            "滚动 RankIC 权重只使用历史信息，无法提前知道未来最强因子，因此会在风格切换中存在滞后。",
            "单因子样本内冠军本身带有事后选择偏差。实盘前并不知道未来哪个单因子会成为冠军，多因子更接近真实决策过程。",
        ]
    )

    add_heading(doc, "7. 多因子的价值", 1)
    add_callout(
        doc,
        "多因子的核心价值",
        "多因子的目标不是保证超过样本期最强单因子，而是在未来风格不确定时降低单一因子失效风险，并提供更稳健、可解释、可扩展的组合信号框架。",
        fill=LIGHT_BLUE_FILL,
    )
    add_bullets(
        doc,
        [
            "分散单因子失效风险，避免组合过度依赖某一个风格。",
            "通过滚动 RankIC 实现因子轮动，弱因子失效时权重会自然下降。",
            "在等权合成明显失效的情况下，RankIC 加权仍能产生正向 D10-D1 收益，说明模型具有实际筛选能力。",
            "为后续策略优化提供底座，可继续叠加精选因子池、约束优化和样本外测试。",
        ]
    )

    add_heading(doc, "8. 权重结构与金融解释", 1)
    active_weights = (
        weights[(weights["composite_factor"].eq("Composite_RollingRankIC")) & (weights["weight_source"].eq("rolling_rankic"))]
        .groupby("source_factor")["weight"]
        .mean()
        .sort_values(ascending=False)
        .head(8)
    )
    weight_rows = [[idx, pct(value)] for idx, value in active_weights.items()]
    add_table(doc, ["源因子", "平均权重"], weight_rows, [2.8, 1.2], numeric_cols={1})
    add_para(
        doc,
        "权重主要集中在 Volatility、Dividend_Yield 和 PE_TTM，和单因子表现较强的方向基本一致，说明 RankIC 加权机制在金融逻辑上是自洽的。",
    )
    doc.add_picture(str(IMAGES / "rankic_weights.png"), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "图 2：滚动 RankIC 多因子的源因子权重变化。", size=9.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "9. 风险与改进方向", 1)
    add_callout(
        doc,
        "主要风险",
        "当前样本只有 54 个调仓期，且分层收益并非严格单调。结果可以支持继续研究，但不应直接视为可上线策略。",
        fill=RISK_FILL,
    )
    add_bullets(
        doc,
        [
            "样本期较短，建议扩展至更长历史区间，并做分年度检验。",
            "当前纳入全部 11 个因子，建议筛选长期 RankIC 为正且经济含义稳定的因子池。",
            "交易成本只使用 10bp 单边假设，后续应加入冲击成本、涨跌停成交约束和容量评估。",
            "需要做滚动样本外测试，用过去数据选择单因子冠军，再与多因子比较，避免事后选择偏差。",
        ]
    )

    add_heading(doc, "10. 给导师的阶段性总结", 1)
    add_para(
        doc,
        "本阶段完成了从单因子检测到多因子合成的研究闭环。结果显示，简单等权并不可取，而滚动 RankIC 加权能够显著改善多因子排序能力。虽然当前多因子没有超过样本期最强单因子，但其优势体现在降低单一风格依赖、提供动态因子轮动机制和形成更稳健的研究框架。",
    )
    add_para(
        doc,
        "下一阶段建议重点推进精选因子池和滚动样本外测试。如果精选后的多因子能够在更多年份、更多股票池和更高交易成本下保持稳定表现，其策略价值会比当前全量因子版本更明确。",
    )

    doc.core_properties.title = "A股多因子合成与分层回测研究"
    doc.core_properties.subject = "多因子研究报告"
    doc.core_properties.author = "share_quant research"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
