"""Run the frozen GroupBalanced11_v1 extended-history validation.

This script intentionally lives outside the existing research package so the
historical implementation remains untouched.  It reuses the existing factor,
neutralisation, execution, and backtest primitives and writes only beneath its
own output directory.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Sequence

# Make rankings, regressions, and floating reductions reproducible across runs.
for _thread_variable in (
    "OMP_NUM_THREADS",
    "OPENBLAS_NUM_THREADS",
    "MKL_NUM_THREADS",
    "NUMEXPR_NUM_THREADS",
):
    os.environ[_thread_variable] = "1"

import duckdb
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from factor_research.factor_backtest import (
    FactorResearchRunner,
    FactorSpec,
    backtest_monthly_top_quantile,
    build_rebalance_calendar,
    filter_scores_for_execution_universe,
    neutralize_scores,
    normalize_factor_panel,
    select_top_quantile_weights,
)
from factor_research.factor_diagnostics import (
    assign_score_deciles,
    build_forward_period_returns,
    compute_decile_returns,
    compute_ic_by_period,
    compute_long_short_returns,
)
from factor_research.factor_marginal_analysis import (
    FIXED_4,
    FULL_11,
    GROUPED_8,
    OOS_FULL,
    PeriodSpec,
    _coverage_diagnostics,
    _git_status,
    _sha256,
    build_all_composite_scores,
    build_bucket_turnover,
    build_experiment_definitions,
    build_monthly_strategy_results,
    compare_protected_snapshots,
    factor_group_definition,
    load_analysis_config,
    snapshot_protected_files,
    summarize_strategy_metrics,
)
from factor_research.paths import load_paths


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = PROJECT_ROOT.parent
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "outputs" / "group_balanced11_extended_validation"
MARGINAL_CONFIG = PROJECT_ROOT / "configs" / "factor_marginal_analysis.yaml"
OLD_METRICS = PROJECT_ROOT / "outputs" / "factor_marginal_analysis" / "grouped_composite_metrics.csv"

WARMUP_START = "20150101"
SIGNAL_START = "20160101"
REFERENCE_START = "20250101"
REFERENCE_END = "20260707"
CHALLENGER = "GroupBalanced11_v1"
BOOTSTRAP_SEED = 20260722
BOOTSTRAP_REPETITIONS = 10_000
BOOTSTRAP_BLOCK_MONTHS = 6

INTERNAL_STRATEGIES = (FULL_11, GROUPED_8, FIXED_4)
STRATEGY_NAMES = {
    FULL_11: FULL_11,
    GROUPED_8: CHALLENGER,
    FIXED_4: FIXED_4,
}
DISPLAY_NAMES = {
    FULL_11: "原始11因子等权",
    CHALLENGER: "GroupBalanced11_v1",
    FIXED_4: "固定4因子等权",
}
PRIMARY_METRICS = {
    "d10_d1": "d10_d1_annual_return",
    "top20_excess": "top20_excess_annual_return",
}
REFERENCE_COLUMNS = (
    "d10_d1_annual_return",
    "d10_d1_cumulative_return",
    "d10_d1_max_drawdown",
    "d10_d1_average_two_leg_turnover",
    "top20_annual_return",
    "benchmark_annual_return",
    "top20_excess_annual_return",
    "top20_cumulative_return",
    "top20_max_drawdown",
    "top20_average_monthly_turnover",
    "rank_ic_mean",
    "average_valid_stock_count",
    "average_valid_stock_coverage",
    "average_available_component_count",
)


@dataclass
class PreparedResearch:
    db_path: Path
    end_date: str
    evaluation_start: str
    calendar: pd.DataFrame
    returns: pd.DataFrame
    suspensions: pd.DataFrame
    benchmark: pd.DataFrame
    execution_universe: pd.DataFrame
    composite_scores: pd.DataFrame
    group_scores: pd.DataFrame
    quality_rows: list[dict[str, object]]


@dataclass
class BacktestArtifacts:
    calendar: pd.DataFrame
    composite_scores: pd.DataFrame
    long_short: pd.DataFrame
    rank_ic: pd.DataFrame
    two_leg_turnover: pd.DataFrame
    top20_daily: pd.DataFrame
    top20_turnover: pd.DataFrame


def parse_args() -> argparse.Namespace:
    paths = load_paths(None)
    parser = argparse.ArgumentParser(
        description="Validate frozen GroupBalanced11_v1 on the extended 2016-present history."
    )
    parser.add_argument("--database-path", type=Path, default=paths.database_path)
    parser.add_argument("--end", default=None, help="Optional end date, YYYY-MM-DD.")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    return parser.parse_args()


def compact_date(value: str | None) -> str | None:
    return value.replace("-", "") if value else None


def dashed_date(value: str) -> str:
    return f"{value[:4]}-{value[4:6]}-{value[6:8]}"


def add_quality(
    rows: list[dict[str, object]],
    check: str,
    status: str,
    detail: str,
    *,
    critical: bool,
) -> None:
    rows.append(
        {
            "check": check,
            "status": status,
            "critical": critical,
            "detail": detail,
        }
    )


def latest_quality_status(con: duckdb.DuckDBPyConnection) -> pd.DataFrame:
    return con.execute(
        """
        select dataset, check_name, status, detail, checked_at
        from (
            select *, row_number() over (
                partition by dataset, check_name order by checked_at desc
            ) as rn
            from data_quality
        )
        where rn = 1
          and (
              (dataset = '__database__' and check_name in (
                  'cross:daily_adj_factor',
                  'cross:daily_trade_calendar',
                  'view:v_adjusted_returns',
                  'view:v_stock_universe_daily',
                  'view:v_fina_indicator_asof_intervals',
                  'view:v_index_data'
              ))
              or (dataset = 'fina_indicator_vip' and check_name = 'dataset_integrity')
          )
        order by dataset, check_name
        """
    ).fetchdf()


def validate_frozen_definition(config: object) -> None:
    expected_groups = {
        "Value": ("PE_TTM", "Dividend_Yield"),
        "Quality": ("ROE", "Debt_to_Equity", "Gross_Margin"),
        "Growth": ("Revenue_Growth",),
        "LowVol": ("Volatility",),
        "Momentum": ("Momentum_60D",),
        "FundFlow": ("Main_Net_In",),
        "Liquidity": ("Turnover_20D",),
        "Ownership": ("Holder_Concen",),
    }
    actual_groups = {group.group_id: group.factor_ids for group in config.groups}
    if len(config.factor_ids) != 11 or actual_groups != expected_groups:
        raise RuntimeError("the frozen 11-factor/8-group definition has drifted")
    if config.minimum_component_count(11) != 6 or config.minimum_component_count(8) != 5:
        raise RuntimeError("the frozen 6/11 or 5/8 coverage threshold has drifted")


def prepare_research(
    db_path: Path,
    requested_end: str | None,
    output_dir: Path,
) -> tuple[PreparedResearch, object]:
    config = load_analysis_config(MARGINAL_CONFIG)
    validate_frozen_definition(config)
    quality: list[dict[str, object]] = []
    runner = FactorResearchRunner(
        project_root=WORKSPACE_ROOT / "database",
        db_path=db_path,
        output_dir=output_dir,
        start="2016-01-01",
        end=dashed_date(requested_end) if requested_end else None,
        warmup_start="2015-01-01",
        benchmark=config.benchmark,
        transaction_cost=config.transaction_cost,
    )

    with duckdb.connect(str(db_path), read_only=True) as con:
        con.execute("set threads=1")
        latest = runner._latest_trade_date(con)
        end_date = requested_end or latest
        if end_date > latest:
            raise ValueError(f"requested end {end_date} exceeds database end {latest}")

        adjusted = con.execute(
            "select min(trade_date), max(trade_date), count(*) from v_adjusted_returns"
        ).fetchone()
        adjusted_ok = bool(adjusted and str(adjusted[0]) <= "20150105" and str(adjusted[1]) >= end_date)
        add_quality(
            quality,
            "adjusted_return_coverage",
            "passed" if adjusted_ok else "failed",
            f"min={adjusted[0]}, max={adjusted[1]}, rows={adjusted[2]}",
            critical=True,
        )

        benchmark_row = con.execute(
            """
            select min(trade_date), max(trade_date), count(*)
            from v_index_data
            where ts_code = ? and trade_date between ? and ? and close is not null
            """,
            [config.benchmark, WARMUP_START, end_date],
        ).fetchone()
        benchmark_ok = bool(
            benchmark_row
            and benchmark_row[0]
            and str(benchmark_row[0]) <= "20150105"
            and str(benchmark_row[1]) >= end_date
        )
        add_quality(
            quality,
            "benchmark_coverage",
            "passed" if benchmark_ok else "failed",
            f"benchmark={config.benchmark}, min={benchmark_row[0]}, max={benchmark_row[1]}, rows={benchmark_row[2]}",
            critical=True,
        )

        for row in latest_quality_status(con).to_dict("records"):
            is_financial_null_warning = (
                row["dataset"] == "fina_indicator_vip"
                and row["check_name"] == "dataset_integrity"
                and row["status"] != "passed"
            )
            add_quality(
                quality,
                f"catalog:{row['dataset']}:{row['check_name']}",
                "warning" if is_financial_null_warning else str(row["status"]),
                f"{row['detail']} (checked_at={row['checked_at']})"
                + (
                    "; rows without announcement visibility are excluded by the existing factor loader"
                    if is_financial_null_warning
                    else ""
                ),
                critical=not is_financial_null_warning,
            )

        all_trade_dates = runner._load_trade_dates(con, WARMUP_START, end_date)
        calendar = build_rebalance_calendar(all_trade_dates, SIGNAL_START, end_date)
        if calendar.empty:
            raise RuntimeError("no rebalance calendar is available")
        evaluation_start = str(calendar["exec_date"].min())
        signal_dates = calendar["signal_date"].drop_duplicates().reset_index(drop=True)
        con.register("signal_dates", pd.DataFrame({"trade_date": signal_dates}))
        raw = runner._load_factor_panel(con, signal_dates, end_date)
        execution_universe = runner._load_execution_universe(con, calendar)
        exposures = runner._load_exposures(con, signal_dates)
        returns = runner._load_returns(con, SIGNAL_START, end_date)
        suspensions = runner._load_suspensions(con, SIGNAL_START, end_date)
        benchmark = runner._load_benchmark(con, WARMUP_START, end_date)

    first_by_factor = raw.groupby("factor")["trade_date"].min().to_dict()
    factors_ok = set(first_by_factor) == set(config.factor_ids)
    volatility_ok = first_by_factor.get("Volatility", "99999999") <= "20160131"
    add_quality(
        quality,
        "frozen_factor_availability",
        "passed" if factors_ok and volatility_ok else "failed",
        f"factor_count={len(first_by_factor)}, first_volatility={first_by_factor.get('Volatility')}",
        critical=True,
    )
    add_quality(
        quality,
        "evaluation_boundary",
        "passed" if evaluation_start.startswith("201602") else "failed",
        f"first_signal={calendar['signal_date'].min()}, first_execution={evaluation_start}",
        critical=True,
    )

    specs = tuple(
        FactorSpec(
            name=factor,
            category="group_balanced11_extended_validation",
            chinese_name=factor,
            formula="loaded_from_existing_factor_implementation",
            direction=direction,
        )
        for factor, direction in config.factor_directions.items()
    )
    scores = normalize_factor_panel(raw[raw["factor"].isin(config.factor_ids)], specs=specs)
    neutralized = neutralize_scores(scores, exposures)
    tradable = filter_scores_for_execution_universe(neutralized, calendar, execution_universe)

    definitions = [
        definition
        for definition in build_experiment_definitions(config)
        if definition.strategy_id in {FULL_11, FIXED_4}
    ]
    composite_scores, group_scores = build_all_composite_scores(
        tradable,
        definitions,
        config.groups,
        GROUPED_8,
        config.minimum_component_count(len(config.groups)),
        config.factor_ids,
    )
    duplicate_count = int(
        composite_scores.duplicated(["factor", "trade_date", "ts_code"]).sum()
    )
    strategy_signal_counts = (
        composite_scores.groupby("factor")["trade_date"].nunique().reindex(INTERNAL_STRATEGIES)
    )
    expected_signals = int(calendar["signal_date"].nunique())
    calendar_ok = bool(
        duplicate_count == 0
        and strategy_signal_counts.notna().all()
        and (strategy_signal_counts == expected_signals).all()
    )
    add_quality(
        quality,
        "shared_strategy_calendar",
        "passed" if calendar_ok else "failed",
        f"expected_signals={expected_signals}, strategy_signals={strategy_signal_counts.to_dict()}, duplicates={duplicate_count}",
        critical=True,
    )

    return (
        PreparedResearch(
            db_path=db_path,
            end_date=end_date,
            evaluation_start=evaluation_start,
            calendar=calendar,
            returns=returns,
            suspensions=suspensions,
            benchmark=benchmark,
            execution_universe=execution_universe,
            composite_scores=composite_scores,
            group_scores=group_scores,
            quality_rows=quality,
        ),
        config,
    )


def run_backtest(prepared: PreparedResearch, config: object, *, end_date: str) -> BacktestArtifacts:
    calendar = prepared.calendar[
        prepared.calendar["exec_date"].astype(str) <= end_date
    ].copy()
    signals = set(calendar["signal_date"].astype(str))
    composite = prepared.composite_scores[
        prepared.composite_scores["trade_date"].astype(str).isin(signals)
    ].copy()
    returns = prepared.returns[prepared.returns["trade_date"].astype(str) <= end_date].copy()
    suspensions = prepared.suspensions[
        prepared.suspensions["trade_date"].astype(str) <= end_date
    ].copy()

    forward = build_forward_period_returns(
        returns,
        calendar,
        end_date,
        ts_codes=composite["ts_code"].dropna().astype(str).unique(),
        suspensions=suspensions,
    )
    deciles = assign_score_deciles(composite, bucket_count=config.bucket_count, score_col="score")
    layer_returns = compute_decile_returns(deciles, forward)
    long_short = compute_long_short_returns(layer_returns, bucket_count=config.bucket_count)
    rank_ic = compute_ic_by_period(composite, forward, score_col="score")
    two_leg_turnover = build_bucket_turnover(deciles, bucket_count=config.bucket_count)
    weights = select_top_quantile_weights(
        composite, quantile=config.top_quantile, score_col="score"
    )
    top20_daily, top20_turnover = backtest_monthly_top_quantile(
        weights,
        returns,
        calendar,
        end_date=end_date,
        transaction_cost=config.transaction_cost,
        suspensions=suspensions,
    )
    benchmark = prepared.benchmark[
        prepared.benchmark["trade_date"].astype(str) <= end_date
    ][["trade_date", "benchmark_return"]].copy()
    benchmark["trade_date"] = benchmark["trade_date"].astype(str)
    top20_daily["trade_date"] = top20_daily["trade_date"].astype(str)
    top20_daily = top20_daily.merge(benchmark, on="trade_date", how="left")
    missing_benchmark = int(top20_daily["benchmark_return"].isna().sum())
    add_quality(
        prepared.quality_rows,
        f"benchmark_alignment_through_{end_date}",
        "passed" if missing_benchmark == 0 else "failed",
        f"missing_strategy_day_benchmark_rows={missing_benchmark}",
        critical=True,
    )
    top20_daily["benchmark_return"] = pd.to_numeric(
        top20_daily["benchmark_return"], errors="coerce"
    ).fillna(0.0)
    top20_daily["excess_return"] = (
        top20_daily["portfolio_return"] - top20_daily["benchmark_return"]
    )
    return BacktestArtifacts(
        calendar=calendar,
        composite_scores=composite,
        long_short=long_short,
        rank_ic=rank_ic,
        two_leg_turnover=two_leg_turnover,
        top20_daily=top20_daily,
        top20_turnover=top20_turnover,
    )


def period_specs(evaluation_start: str, end_date: str) -> list[PeriodSpec]:
    first_year = int(evaluation_start[:4])
    end_year = int(end_date[:4])
    periods = [
        PeriodSpec(
            str(year),
            evaluation_start if year == first_year else f"{year}0101",
            f"{year}1231",
        )
        for year in range(first_year, end_year + 1)
    ]
    if end_date >= REFERENCE_START:
        periods.append(PeriodSpec("REFERENCE_2025_20260707", REFERENCE_START, REFERENCE_END))
    return periods


def summarize(
    prepared: PreparedResearch,
    artifacts: BacktestArtifacts,
    end_date: str,
    periods: Sequence[PeriodSpec],
    evaluation_start: str | None = None,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    start = evaluation_start or prepared.evaluation_start
    metrics = summarize_strategy_metrics(
        strategy_ids=INTERNAL_STRATEGIES,
        periods=periods,
        end_date=end_date,
        long_short_returns=artifacts.long_short,
        rank_ic=artifacts.rank_ic,
        top20_daily=artifacts.top20_daily,
        top20_turnover=artifacts.top20_turnover,
        two_leg_turnover=artifacts.two_leg_turnover,
        composite_scores=artifacts.composite_scores,
        calendar=artifacts.calendar,
        execution_universe=prepared.execution_universe,
        oos_start=start,
    )
    monthly = build_monthly_strategy_results(
        INTERNAL_STRATEGIES,
        artifacts.long_short,
        artifacts.rank_ic,
        artifacts.top20_daily,
        artifacts.top20_turnover,
        artifacts.two_leg_turnover,
        artifacts.calendar,
        start,
        end_date,
    )
    return metrics, monthly


def complete_years(evaluation_start: str, end_date: str) -> list[int]:
    first = int(evaluation_start[:4]) + 1
    last = int(end_date[:4]) - 1
    years = list(range(first, last + 1)) if last >= first else []
    if end_date[4:6] == "12" and int(end_date[6:8]) >= 20:
        years.append(int(end_date[:4]))
    return sorted(set(years))


def strategy_name(value: object) -> object:
    return STRATEGY_NAMES.get(str(value), value)


def rename_strategies(frame: pd.DataFrame) -> pd.DataFrame:
    result = frame.copy()
    for column in ("strategy_id", "factor", "included_strategy", "excluded_strategy"):
        if column in result.columns:
            result[column] = result[column].map(strategy_name)
    return result


def monthly_advantage(monthly: pd.DataFrame) -> pd.DataFrame:
    columns = {
        "d10_d1_return": "d10_d1_delta",
        "top20_excess_return": "top20_excess_delta",
    }
    pieces: list[pd.Series] = []
    for source, target in columns.items():
        pivot = monthly.pivot(index="attribution_month", columns="factor", values=source)
        pieces.append((pivot[GROUPED_8] - pivot[FULL_11]).rename(target))
    result = pd.concat(pieces, axis=1).reset_index()
    result["year"] = result["attribution_month"].str[:4].astype(int)
    return result


def newey_west_mean(values: np.ndarray) -> tuple[float, float, int]:
    n = int(values.size)
    if n < 2:
        return math.nan, math.nan, 0
    demeaned = values - values.mean()
    lag = max(1, int(math.floor(4.0 * (n / 100.0) ** (2.0 / 9.0))))
    lag = min(lag, n - 1)
    long_variance = float(np.dot(demeaned, demeaned) / n)
    for offset in range(1, lag + 1):
        covariance = float(np.dot(demeaned[offset:], demeaned[:-offset]) / n)
        long_variance += 2.0 * (1.0 - offset / (lag + 1.0)) * covariance
    se = math.sqrt(max(long_variance, 0.0) / n)
    statistic = float(values.mean() / se) if se > 0 else math.nan
    p_value = math.erfc(abs(statistic) / math.sqrt(2.0)) if math.isfinite(statistic) else math.nan
    return statistic, p_value, lag


def moving_block_bootstrap(
    values: np.ndarray,
    *,
    repetitions: int,
    block_length: int,
    seed: int,
) -> tuple[float, float]:
    n = int(values.size)
    if n == 0:
        return math.nan, math.nan
    block = min(block_length, n)
    starts = np.arange(0, n - block + 1)
    blocks = np.stack([values[start : start + block] for start in starts])
    draws_per_sample = math.ceil(n / block)
    rng = np.random.default_rng(seed)
    means = np.empty(repetitions, dtype=float)
    for idx in range(repetitions):
        chosen = rng.integers(0, len(blocks), size=draws_per_sample)
        means[idx] = blocks[chosen].reshape(-1)[:n].mean()
    low, high = np.quantile(means, [0.025, 0.975])
    return float(low), float(high)


def statistical_tests(advantage: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for index, (metric, column) in enumerate(
        (("d10_d1", "d10_d1_delta"), ("top20_excess", "top20_excess_delta"))
    ):
        values = pd.to_numeric(advantage[column], errors="coerce").dropna().to_numpy(dtype=float)
        n = int(values.size)
        std = float(values.std(ddof=1)) if n > 1 else math.nan
        simple_t = float(values.mean() / (std / math.sqrt(n))) if n > 1 and std > 0 else math.nan
        hac_stat, hac_p, hac_lag = newey_west_mean(values)
        ci_low, ci_high = moving_block_bootstrap(
            values,
            repetitions=BOOTSTRAP_REPETITIONS,
            block_length=BOOTSTRAP_BLOCK_MONTHS,
            seed=BOOTSTRAP_SEED + index,
        )
        rows.append(
            {
                "metric": metric,
                "months": n,
                "mean_delta": float(values.mean()) if n else math.nan,
                "median_delta": float(np.median(values)) if n else math.nan,
                "win_rate": float((values > 0).mean()) if n else math.nan,
                "simple_t": simple_t,
                "newey_west_lag": hac_lag,
                "newey_west_stat": hac_stat,
                "newey_west_two_sided_p": hac_p,
                "bootstrap_repetitions": BOOTSTRAP_REPETITIONS,
                "bootstrap_block_months": min(BOOTSTRAP_BLOCK_MONTHS, n),
                "bootstrap_ci_95_low": ci_low,
                "bootstrap_ci_95_high": ci_high,
                "bootstrap_significant_positive": bool(ci_low > 0),
            }
        )
    return pd.DataFrame(rows)


def annual_stability(metrics: pd.DataFrame, years: Sequence[int]) -> pd.DataFrame:
    indexed = metrics.set_index(["period", "strategy_id"])
    rows: list[dict[str, object]] = []
    for year in years:
        period = str(year)
        if (period, GROUPED_8) not in indexed.index or (period, FULL_11) not in indexed.index:
            continue
        for name, metric in PRIMARY_METRICS.items():
            challenger_value = float(indexed.loc[(period, GROUPED_8), metric])
            baseline_value = float(indexed.loc[(period, FULL_11), metric])
            rows.append(
                {
                    "year": year,
                    "metric": name,
                    "challenger": challenger_value,
                    "baseline": baseline_value,
                    "delta": challenger_value - baseline_value,
                    "challenger_wins": bool(challenger_value > baseline_value),
                }
            )
    return pd.DataFrame(rows)


def build_decision(
    metrics: pd.DataFrame,
    stability: pd.DataFrame,
    tests: pd.DataFrame,
    years: Sequence[int],
    quality: pd.DataFrame,
) -> dict[str, object]:
    full = metrics[metrics["period"] == OOS_FULL].set_index("strategy_id")
    challenger = full.loc[GROUPED_8]
    baseline = full.loc[FULL_11]
    threshold = math.ceil(2 * len(years) / 3) if years else 0
    win_counts = (
        stability.groupby("metric")["challenger_wins"].sum().astype(int).to_dict()
        if not stability.empty
        else {}
    )
    returns = {
        key: bool(challenger[metric] > baseline[metric])
        for key, metric in PRIMARY_METRICS.items()
    }
    stable = {
        key: bool(win_counts.get(key, 0) >= threshold and threshold > 0)
        for key in PRIMARY_METRICS
    }
    significant = {
        str(row["metric"]): bool(row["bootstrap_significant_positive"])
        for row in tests.to_dict("records")
    }
    drawdown = {
        "d10_d1": bool(challenger["d10_d1_max_drawdown"] >= baseline["d10_d1_max_drawdown"]),
        "top20": bool(challenger["top20_max_drawdown"] >= baseline["top20_max_drawdown"]),
    }
    turnover = {
        "d10_d1": bool(
            challenger["d10_d1_average_two_leg_turnover"]
            <= baseline["d10_d1_average_two_leg_turnover"]
        ),
        "top20": bool(
            challenger["top20_average_monthly_turnover"]
            <= baseline["top20_average_monthly_turnover"]
        ),
    }
    critical_failures = quality[(quality["critical"] == True) & (quality["status"] != "passed")]
    performance_gate = all(returns.values()) and all(stable.values())
    strict_gate = (
        performance_gate
        and all(significant.get(metric, False) for metric in PRIMARY_METRICS)
        and all(drawdown.values())
        and all(turnover.values())
    )
    if not critical_failures.empty:
        label = "数据不足，无法判定"
    elif strict_gate:
        label = "可替换/挑战者更好"
    elif performance_gate:
        label = "继续作为挑战者"
    else:
        label = "不优于基准"
    return {
        "label": label,
        "evaluation_period": {
            "start": str(metrics.loc[metrics["period"] == OOS_FULL, "period_start"].iloc[0]),
            "end": str(metrics.loc[metrics["period"] == OOS_FULL, "period_end"].iloc[0]),
        },
        "complete_years": list(map(int, years)),
        "required_winning_years": threshold,
        "winning_years": win_counts,
        "full_period_return_gates": returns,
        "annual_stability_gates": stable,
        "bootstrap_significance_gates": significant,
        "drawdown_gates": drawdown,
        "turnover_gates": turnover,
        "critical_quality_failures": critical_failures["check"].astype(str).tolist(),
    }


def reference_reproduction(
    prepared: PreparedResearch,
    config: object,
) -> tuple[pd.DataFrame, bool]:
    if prepared.end_date < REFERENCE_END or not OLD_METRICS.exists():
        return pd.DataFrame(), True
    artifacts = run_backtest(prepared, config, end_date=REFERENCE_END)
    metrics, _monthly = summarize(
        prepared,
        artifacts,
        REFERENCE_END,
        periods=[],
        evaluation_start=REFERENCE_START,
    )
    new = metrics[metrics["period"] == OOS_FULL].set_index("strategy_id")
    old = pd.read_csv(OLD_METRICS)
    old = old[old["period"] == OOS_FULL].set_index("strategy_id")
    rows: list[dict[str, object]] = []
    passed = True
    for strategy in INTERNAL_STRATEGIES:
        for metric in REFERENCE_COLUMNS:
            new_value = pd.to_numeric(pd.Series([new.loc[strategy, metric]]), errors="coerce").iloc[0]
            old_value = pd.to_numeric(pd.Series([old.loc[strategy, metric]]), errors="coerce").iloc[0]
            equal = bool(
                (pd.isna(new_value) and pd.isna(old_value))
                or np.isclose(new_value, old_value, rtol=1e-9, atol=1e-10, equal_nan=True)
            )
            passed = passed and equal
            rows.append(
                {
                    "strategy_id": STRATEGY_NAMES[strategy],
                    "metric": metric,
                    "new_value": new_value,
                    "old_value": old_value,
                    "absolute_difference": (
                        abs(float(new_value) - float(old_value))
                        if pd.notna(new_value) and pd.notna(old_value)
                        else math.nan
                    ),
                    "within_tolerance": equal,
                }
            )
    return pd.DataFrame(rows), passed


def markdown_table(frame: pd.DataFrame) -> str:
    if frame.empty:
        return "（无数据）"
    display = frame.fillna("—").astype(str)
    headers = list(display.columns)
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in display.to_numpy().tolist())
    return "\n".join(lines)


def pct(value: object) -> str:
    number = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
    return "—" if pd.isna(number) else f"{float(number):.2%}"


def num(value: object, digits: int = 3) -> str:
    number = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
    return "—" if pd.isna(number) else f"{float(number):.{digits}f}"


def build_report_markdown(
    metrics: pd.DataFrame,
    stability: pd.DataFrame,
    tests: pd.DataFrame,
    quality: pd.DataFrame,
    decision: dict[str, object],
    reproduction: pd.DataFrame,
) -> str:
    full = rename_strategies(metrics[metrics["period"] == OOS_FULL]).copy()
    full_table = full[
        [
            "strategy_id",
            "d10_d1_annual_return",
            "top20_excess_annual_return",
            "rank_ic_mean",
            "d10_d1_max_drawdown",
            "top20_max_drawdown",
            "d10_d1_average_two_leg_turnover",
            "top20_average_monthly_turnover",
        ]
    ].copy()
    full_table.columns = [
        "方案",
        "D10-D1年化",
        "Top20年化超额",
        "RankIC",
        "D10最大回撤",
        "Top20最大回撤",
        "D10双边换手",
        "Top20换手",
    ]
    for column in ("D10-D1年化", "Top20年化超额", "D10最大回撤", "Top20最大回撤", "D10双边换手", "Top20换手"):
        full_table[column] = full_table[column].map(pct)
    full_table["RankIC"] = full_table["RankIC"].map(lambda value: num(value, 4))

    stable_table = stability.pivot(index="year", columns="metric", values="delta").reset_index()
    if not stable_table.empty:
        stable_table.columns.name = None
        stable_table = stable_table.rename(
            columns={"d10_d1": "D10-D1差", "top20_excess": "Top20超额差"}
        )
        for column in ("D10-D1差", "Top20超额差"):
            stable_table[column] = stable_table[column].map(pct)

    test_table = tests[
        [
            "metric",
            "months",
            "mean_delta",
            "win_rate",
            "simple_t",
            "newey_west_stat",
            "newey_west_two_sided_p",
            "bootstrap_ci_95_low",
            "bootstrap_ci_95_high",
        ]
    ].copy()
    test_table["mean_delta"] = test_table["mean_delta"].map(pct)
    test_table["win_rate"] = test_table["win_rate"].map(pct)
    for column in ("simple_t", "newey_west_stat", "newey_west_two_sided_p"):
        test_table[column] = test_table[column].map(lambda value: num(value, 3))
    test_table["bootstrap_ci_95_low"] = test_table["bootstrap_ci_95_low"].map(pct)
    test_table["bootstrap_ci_95_high"] = test_table["bootstrap_ci_95_high"].map(pct)

    critical = quality[(quality["critical"] == True) & (quality["status"] != "passed")]
    warnings = quality[quality["status"] == "warning"]
    reproduction_text = (
        "未执行（参考产物或参考截止日不可用）。"
        if reproduction.empty
        else f"{int(reproduction['within_tolerance'].sum())}/{len(reproduction)} 个数值在容差内。"
    )
    return f"""# GroupBalanced11_v1 全历史严格验证

## 结论

**{decision['label']}**

- 评价期：{decision['evaluation_period']['start']} 至 {decision['evaluation_period']['end']}。
- 完整年度：{', '.join(map(str, decision['complete_years']))}；每个主指标至少需在 {decision['required_winning_years']} 个年度跑赢。
- D10-D1 胜出年度数：{decision['winning_years'].get('d10_d1', 0)}；Top20 超额胜出年度数：{decision['winning_years'].get('top20_excess', 0)}。
- 本结论是历史回测证据，不等同于方案冻结后的前瞻样本外验证。

## 全期三方案比较

{markdown_table(full_table)}

## 完整年度稳定性

{markdown_table(stable_table)}

## 月度配对统计

{markdown_table(test_table)}

## 严格门禁

```json
{json.dumps(decision, ensure_ascii=False, indent=2)}
```

## 数据质量与复现

- 关键质量失败数：{len(critical)}；警告数：{len(warnings)}。
- 2025-01-01 至 2026-07-07 旧区间复现：{reproduction_text}
- 数据库全程以只读方式打开；因子、方向、分组、执行时点和成本口径均来自现有冻结实现。
"""


def configure_docx(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(9.5)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_docx_table(doc: Document, frame: pd.DataFrame) -> None:
    if frame.empty:
        doc.add_paragraph("无数据")
        return
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(frame.columns):
        table.rows[0].cells[idx].text = str(column)
    for row in frame.fillna("—").astype(str).itertuples(index=False, name=None):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def write_docx_report(
    path: Path,
    metrics: pd.DataFrame,
    stability: pd.DataFrame,
    tests: pd.DataFrame,
    quality: pd.DataFrame,
    decision: dict[str, object],
    image_paths: Sequence[Path],
) -> None:
    doc = Document()
    configure_docx(doc)
    title = doc.add_heading("GroupBalanced11_v1 全历史严格验证", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("结论", level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(decision["label"]))
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph(
        f"评价期 {decision['evaluation_period']['start']} 至 {decision['evaluation_period']['end']}；"
        f"完整年度需至少 {decision['required_winning_years']} 年胜出。"
    )
    doc.add_paragraph("本结论属于历史回测证据，不等同于方案冻结后的前瞻样本外验证。")

    doc.add_heading("全期三方案比较", level=1)
    full = rename_strategies(metrics[metrics["period"] == OOS_FULL])[
        [
            "strategy_id",
            "d10_d1_annual_return",
            "top20_excess_annual_return",
            "d10_d1_max_drawdown",
            "top20_max_drawdown",
            "d10_d1_average_two_leg_turnover",
            "top20_average_monthly_turnover",
        ]
    ].copy()
    full.columns = ["方案", "D10-D1年化", "Top20超额", "D10回撤", "Top20回撤", "D10换手", "Top20换手"]
    for column in full.columns[1:]:
        full[column] = full[column].map(pct)
    add_docx_table(doc, full)

    doc.add_heading("完整年度稳定性", level=1)
    annual = stability.copy()
    if not annual.empty:
        annual["challenger"] = annual["challenger"].map(pct)
        annual["baseline"] = annual["baseline"].map(pct)
        annual["delta"] = annual["delta"].map(pct)
    add_docx_table(doc, annual)

    doc.add_heading("月度配对统计", level=1)
    stats = tests.copy()
    for column in ("mean_delta", "median_delta", "win_rate", "bootstrap_ci_95_low", "bootstrap_ci_95_high"):
        stats[column] = stats[column].map(pct)
    add_docx_table(doc, stats)

    doc.add_heading("严格判定门禁", level=1)
    gates = pd.DataFrame(
        [
            ["全期收益", decision["full_period_return_gates"]],
            ["年度稳定", decision["annual_stability_gates"]],
            ["Bootstrap显著性", decision["bootstrap_significance_gates"]],
            ["最大回撤", decision["drawdown_gates"]],
            ["平均换手", decision["turnover_gates"]],
        ],
        columns=["门禁", "结果"],
    )
    gates["结果"] = gates["结果"].map(lambda value: json.dumps(value, ensure_ascii=False))
    add_docx_table(doc, gates)

    doc.add_heading("图表", level=1)
    for image_path in image_paths:
        doc.add_picture(str(image_path), width=Inches(6.4))
        caption = doc.add_paragraph(image_path.stem)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("数据质量", level=1)
    add_docx_table(doc, quality)
    doc.add_paragraph("数据库以 read_only=True 打开；旧代码、旧配置、数据库与旧输出受运行前后审计保护。")
    doc.save(path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
        document.SaveAs2(str(pdf_path.resolve()), FileFormat=17)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def write_images(
    image_dir: Path,
    monthly: pd.DataFrame,
    advantage: pd.DataFrame,
    stability: pd.DataFrame,
    metrics: pd.DataFrame,
    coverage: pd.DataFrame,
) -> list[Path]:
    image_dir.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
    paths: list[Path] = []

    monthly_plot = rename_strategies(monthly)
    fig, axes = plt.subplots(2, 1, figsize=(11, 8), sharex=True)
    for strategy, group in monthly_plot.groupby("factor"):
        ordered = group.sort_values("attribution_month")
        dates = pd.to_datetime(ordered["attribution_month"])
        d10_nav = (1.0 + pd.to_numeric(ordered["d10_d1_return"], errors="coerce").fillna(0.0)).cumprod()
        excess_nav = (1.0 + pd.to_numeric(ordered["top20_excess_return"], errors="coerce").fillna(0.0)).cumprod()
        axes[0].plot(dates, d10_nav, label=DISPLAY_NAMES.get(str(strategy), str(strategy)))
        axes[1].plot(dates, excess_nav, label=DISPLAY_NAMES.get(str(strategy), str(strategy)))
    axes[0].set_title("D10-D1 月度复利净值（毛收益）")
    axes[1].set_title("Top20 月度超额复利净值（已扣成本）")
    for axis in axes:
        axis.grid(alpha=0.25)
        axis.legend()
    fig.tight_layout()
    path = image_dir / "cumulative_performance.png"
    fig.savefig(path, dpi=170)
    plt.close(fig)
    paths.append(path)

    ordered_adv = advantage.sort_values("attribution_month")
    fig, ax = plt.subplots(figsize=(11, 4.8))
    dates = pd.to_datetime(ordered_adv["attribution_month"])
    ax.plot(dates, ordered_adv["d10_d1_delta"].rolling(12, min_periods=6).mean(), label="D10-D1 12月均值差")
    ax.plot(dates, ordered_adv["top20_excess_delta"].rolling(12, min_periods=6).mean(), label="Top20超额 12月均值差")
    ax.axhline(0.0, color="black", linewidth=0.8)
    ax.set_title("GroupBalanced11_v1 相对 Full11 的滚动月度优势")
    ax.legend()
    ax.grid(alpha=0.25)
    fig.tight_layout()
    path = image_dir / "monthly_paired_advantage.png"
    fig.savefig(path, dpi=170)
    plt.close(fig)
    paths.append(path)

    if not stability.empty:
        pivot = stability.pivot(index="year", columns="metric", values="delta")
        fig, ax = plt.subplots(figsize=(11, 4.8))
        pivot.rename(columns={"d10_d1": "D10-D1", "top20_excess": "Top20超额"}).plot.bar(ax=ax)
        ax.axhline(0.0, color="black", linewidth=0.8)
        ax.set_title("完整年度收益优势")
        ax.set_ylabel("年化收益差")
        ax.grid(axis="y", alpha=0.25)
        fig.tight_layout()
        path = image_dir / "annual_advantage.png"
        fig.savefig(path, dpi=170)
        plt.close(fig)
        paths.append(path)

    full = rename_strategies(metrics[metrics["period"] == OOS_FULL]).set_index("strategy_id")
    risk = full[
        [
            "d10_d1_max_drawdown",
            "top20_max_drawdown",
            "d10_d1_average_two_leg_turnover",
            "top20_average_monthly_turnover",
        ]
    ].rename(
        columns={
            "d10_d1_max_drawdown": "D10最大回撤",
            "top20_max_drawdown": "Top20最大回撤",
            "d10_d1_average_two_leg_turnover": "D10换手",
            "top20_average_monthly_turnover": "Top20换手",
        }
    )
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8))
    risk[["D10最大回撤", "Top20最大回撤"]].plot.bar(ax=axes[0], title="全期最大回撤")
    risk[["D10换手", "Top20换手"]].plot.bar(ax=axes[1], title="全期平均换手")
    for axis in axes:
        axis.grid(axis="y", alpha=0.25)
        axis.tick_params(axis="x", rotation=20)
    fig.tight_layout()
    path = image_dir / "risk_and_turnover.png"
    fig.savefig(path, dpi=170)
    plt.close(fig)
    paths.append(path)

    cov = rename_strategies(coverage)
    cov["month"] = pd.to_datetime(cov["trade_date"].astype(str), format="%Y%m%d")
    fig, ax = plt.subplots(figsize=(11, 4.8))
    for strategy, group in cov.groupby("factor"):
        ax.plot(group["month"], group["valid_stock_count"], label=DISPLAY_NAMES.get(str(strategy), str(strategy)))
    ax.set_title("月末有效股票覆盖")
    ax.set_ylabel("有效股票数")
    ax.legend()
    ax.grid(alpha=0.25)
    fig.tight_layout()
    path = image_dir / "coverage.png"
    fig.savefig(path, dpi=170)
    plt.close(fig)
    paths.append(path)
    return paths


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def json_default(value: object) -> object:
    if isinstance(value, (np.integer,)):
        return int(value)
    if isinstance(value, (np.floating,)):
        return float(value)
    if isinstance(value, (np.bool_,)):
        return bool(value)
    if isinstance(value, Path):
        return str(value)
    raise TypeError(f"cannot serialize {type(value).__name__}")


def write_outputs(
    output_dir: Path,
    prepared: PreparedResearch,
    config: object,
    metrics: pd.DataFrame,
    monthly: pd.DataFrame,
    advantage: pd.DataFrame,
    tests: pd.DataFrame,
    stability: pd.DataFrame,
    quality: pd.DataFrame,
    decision: dict[str, object],
    reproduction: pd.DataFrame,
    git_before: str,
    protected_before: dict[str, tuple[int, int]],
) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    image_dir = output_dir / "images"
    coverage = _coverage_diagnostics(prepared.composite_scores)
    groups = factor_group_definition(config)

    frames = {
        "strategy_period_metrics": rename_strategies(metrics),
        "monthly_strategy_metrics": rename_strategies(monthly),
        "monthly_paired_advantage": advantage,
        "statistical_tests": tests,
        "annual_stability": stability,
        "coverage_diagnostics": rename_strategies(coverage),
        "factor_group_definition": groups,
        "data_quality_gate": quality,
        "reference_reproduction": reproduction,
    }
    paths: dict[str, Path] = {}
    for name, frame in frames.items():
        path = output_dir / f"{name}.csv"
        frame.to_csv(path, index=False, encoding="utf-8-sig")
        paths[name] = path

    decision_path = output_dir / "decision_summary.json"
    decision_path.write_text(
        json.dumps(decision, ensure_ascii=False, indent=2, default=json_default), encoding="utf-8"
    )
    paths["decision_summary"] = decision_path

    image_paths = write_images(image_dir, monthly, advantage, stability, metrics, coverage)
    for path in image_paths:
        paths[path.stem] = path

    markdown = build_report_markdown(metrics, stability, tests, quality, decision, reproduction)
    markdown_path = output_dir / "analysis_report.md"
    markdown_path.write_text(markdown, encoding="utf-8")
    paths["analysis_report_markdown"] = markdown_path

    docx_path = output_dir / "group_balanced11_extended_validation_report.docx"
    write_docx_report(docx_path, metrics, stability, tests, quality, decision, image_paths)
    paths["analysis_report_docx"] = docx_path
    pdf_path = output_dir / "group_balanced11_extended_validation_report.pdf"
    convert_docx_to_pdf(docx_path, pdf_path)
    paths["analysis_report_pdf"] = pdf_path

    protected_after = snapshot_protected_files(WORKSPACE_ROOT, output_dir)
    audit = compare_protected_snapshots(protected_before, protected_after)
    audit_path = output_dir / "protected_file_audit.csv"
    audit.to_csv(audit_path, index=False, encoding="utf-8-sig")
    paths["protected_file_audit"] = audit_path
    if not audit.empty and not audit["unchanged"].all():
        changed = audit.loc[~audit["unchanged"], "path"].astype(str).tolist()
        raise RuntimeError(f"protected files changed during analysis: {changed[:10]}")

    git_after = _git_status(WORKSPACE_ROOT)
    manifest = {
        "artifact_id": "group_balanced11_extended_validation_v1",
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "database_path": str(prepared.db_path),
        "database_open_mode": "read_only",
        "data_end_date": prepared.end_date,
        "evaluation_start": prepared.evaluation_start,
        "warmup_start": WARMUP_START,
        "signal_start": SIGNAL_START,
        "strategies": [STRATEGY_NAMES[value] for value in INTERNAL_STRATEGIES],
        "factor_config": str(config.source_factor_config),
        "factor_config_sha256": _sha256(config.source_factor_config),
        "marginal_config": str(MARGINAL_CONFIG),
        "marginal_config_sha256": _sha256(MARGINAL_CONFIG),
        "script": str(Path(__file__).resolve()),
        "script_sha256": file_sha256(Path(__file__).resolve()),
        "bootstrap_seed": BOOTSTRAP_SEED,
        "bootstrap_repetitions": BOOTSTRAP_REPETITIONS,
        "bootstrap_block_months": BOOTSTRAP_BLOCK_MONTHS,
        "deterministic_execution": {
            "duckdb_threads": 1,
            "omp_threads": 1,
            "openblas_threads": 1,
            "mkl_threads": 1,
            "numexpr_threads": 1,
        },
        "decision": decision,
        "reference_reproduction_passed": bool(
            reproduction.empty or reproduction["within_tolerance"].all()
        ),
        "protected_file_audit_passed": bool(audit.empty or audit["unchanged"].all()),
        "protected_file_count": int(len(audit)),
        "git_status_before_run": git_before.splitlines(),
        "git_status_after_run": git_after.splitlines(),
        "outputs": {key: str(value) for key, value in sorted(paths.items())},
    }
    manifest_path = output_dir / "run_manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2, default=json_default), encoding="utf-8"
    )
    paths["run_manifest"] = manifest_path
    return paths


def main() -> int:
    args = parse_args()
    db_path = args.database_path.resolve()
    output_dir = args.output_dir.resolve()
    if output_dir == PROJECT_ROOT / "outputs" / "factor_marginal_analysis":
        raise ValueError("refusing to overwrite the existing factor_marginal_analysis output")
    requested_end = compact_date(args.end)
    protected_before = snapshot_protected_files(WORKSPACE_ROOT, output_dir)
    git_before = _git_status(WORKSPACE_ROOT)

    prepared, config = prepare_research(db_path, requested_end, output_dir)
    artifacts = run_backtest(prepared, config, end_date=prepared.end_date)
    periods = period_specs(prepared.evaluation_start, prepared.end_date)
    metrics, monthly = summarize(prepared, artifacts, prepared.end_date, periods)
    advantage = monthly_advantage(monthly)
    tests = statistical_tests(advantage)
    years = complete_years(prepared.evaluation_start, prepared.end_date)
    stability = annual_stability(metrics, years)

    reproduction, reproduction_passed = reference_reproduction(prepared, config)
    add_quality(
        prepared.quality_rows,
        "reference_2025_to_20260707_reproduction",
        "passed" if reproduction_passed else "warning",
        (
            "reference not available; check skipped"
            if reproduction.empty
            else (
                f"within_tolerance={int(reproduction['within_tolerance'].sum())}/{len(reproduction)}; "
                "the archived CSV was produced from an earlier database snapshot and is retained as a drift diagnostic"
            )
        ),
        critical=False,
    )
    quality = pd.DataFrame(prepared.quality_rows)
    decision = build_decision(metrics, stability, tests, years, quality)
    paths = write_outputs(
        output_dir,
        prepared,
        config,
        metrics,
        monthly,
        advantage,
        tests,
        stability,
        quality,
        decision,
        reproduction,
        git_before,
        protected_before,
    )
    print(f"Decision: {decision['label']}")
    for name, path in sorted(paths.items()):
        print(f"{name}: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
